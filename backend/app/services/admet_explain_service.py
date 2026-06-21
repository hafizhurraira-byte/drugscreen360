import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import HTTPException
from rdkit import Chem
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.database import get_connection, init_db
from app.models.admet_explain_models import (
    AdmetDescriptorExplanation,
    AdmetExplanationReportCreateRequest,
    AdmetExplanationReportCreateResponse,
    AdmetImportantFeature,
    AdmetPredictionExplainRequest,
    AdmetPredictionExplanationResponse,
)
from app.models.project_workspace_models import ProjectAttachRequest
from app.services.admet_domain_service import evaluate_domain_internal, get_or_calculate_domain_stats
from app.services.admet_trained_model_service import (
    FEATURE_COLUMNS,
    discover_trained_models,
    get_active_trained_model_info,
    predict_trained_model,
    validate_trained_model,
)
from app.services.admet_validation_service import get_latest_external_validation_by_model
from app.services.descriptors import calculate_descriptors, parse_smiles
from app.services.project_workspace_service import attach_project_item

EXPLANATION_REPORT_DIR = Path(__file__).resolve().parents[2] / "admet_explanation_reports"
SCIENTIFIC_NOTICE = "Computational explanation only. Requires experimental and external validation."
LIMITATIONS = [
    "This explanation is computational and model-specific only.",
    "Feature importance and coefficients are model diagnostics, not biological causality.",
    "Applicability domain and uncertainty checks are heuristic and dataset-dependent.",
    "External validation and laboratory testing are required before scientific use.",
    "No clinical safety, efficacy, regulatory approval, or market readiness is implied.",
]


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _resolve_model_id(model_id: str | None) -> str:
    if model_id:
        return model_id
    active = get_active_trained_model_info()
    if active.get("status") != "available" or not active.get("model_id"):
        raise HTTPException(status_code=400, detail=f"No active trained ADMET model is available. Active model status: {active.get('status', 'unknown')}.")
    return str(active["model_id"])


def _model_summary(model_id: str) -> dict[str, Any]:
    models = discover_trained_models()
    summary = next((item for item in models if item["model_id"] == model_id), None)
    if not summary:
        raise HTTPException(status_code=404, detail=f"Trained model '{model_id}' not found.")
    validation = validate_trained_model(model_id)
    if not validation["valid"]:
        raise HTTPException(status_code=400, detail=f"Model '{model_id}' is not valid: {', '.join(validation['errors'])}")
    return summary


def _read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _descriptor_features(smiles: str) -> dict[str, float | int | None]:
    desc = calculate_descriptors(smiles).model_dump()
    return {
        "molecular_weight": desc.get("molecular_weight"),
        "logp": desc.get("logp"),
        "tpsa": desc.get("tpsa"),
        "hbd": desc.get("hydrogen_bond_donors"),
        "hba": desc.get("hydrogen_bond_acceptors"),
        "rotatable_bonds": desc.get("rotatable_bonds"),
        "ring_count": desc.get("ring_count"),
        "aromatic_ring_count": desc.get("aromatic_ring_count"),
        "formal_charge": desc.get("formal_charge"),
        "fraction_csp3": desc.get("fraction_csp3"),
    }


def _descriptor_explanations(model_id: str, features: dict[str, float | int | None]) -> tuple[list[AdmetDescriptorExplanation], list[str]]:
    warnings: list[str] = []
    try:
        stats = get_or_calculate_domain_stats(model_id)
    except Exception as exc:
        return [], [f"Training descriptor ranges are not available: {exc}"]

    explanations: list[AdmetDescriptorExplanation] = []
    for index, feature in enumerate(FEATURE_COLUMNS):
        raw_value = features.get(feature)
        if raw_value is None:
            explanations.append(
                AdmetDescriptorExplanation(
                    feature=feature,
                    status="not_available",
                    explanation=f"{feature} could not be calculated for the query molecule.",
                )
            )
            continue
        value = float(raw_value)
        training_min = float(stats["mins"][index])
        training_max = float(stats["maxs"][index])
        training_mean = float(stats["means"][index])
        training_std = float(stats["stds"][index])
        if value > training_max:
            status = "outside_training_range_high"
            explanation = f"{feature} is above the maximum observed in the training data."
        elif value < training_min:
            status = "outside_training_range_low"
            explanation = f"{feature} is below the minimum observed in the training data."
        elif value > training_mean + training_std:
            status = "high_vs_training_mean"
            explanation = f"{feature} is higher than the training-data mean by more than one standard deviation."
        elif value < training_mean - training_std:
            status = "low_vs_training_mean"
            explanation = f"{feature} is lower than the training-data mean by more than one standard deviation."
        else:
            status = "within_training_range"
            explanation = f"{feature} is within the central range of the training descriptors."
        explanations.append(
            AdmetDescriptorExplanation(
                feature=feature,
                query_value=round(value, 6),
                training_min=round(training_min, 6),
                training_max=round(training_max, 6),
                training_mean=round(training_mean, 6),
                training_std=round(training_std, 6),
                status=status,
                explanation=explanation,
            )
        )
    return explanations, warnings


def _load_model_object(model_summary: dict[str, Any]):
    import joblib

    artifact_path = Path(model_summary["artifact_dir"]) / "model.joblib"
    model_data = joblib.load(artifact_path)
    model = model_data.get("model")
    if model is None:
        raise HTTPException(status_code=400, detail="Model artifact did not contain a supported 'model' object.")
    return model


def _important_features(model_summary: dict[str, Any]) -> tuple[list[AdmetImportantFeature], str, list[str]]:
    warnings: list[str] = []
    model = _load_model_object(model_summary)
    source = None
    values = None
    if hasattr(model, "feature_importances_"):
        source = "model_feature_importance"
        values = list(model.feature_importances_)
        summary = "The model exposes global feature_importances_. These are model diagnostics only and do not prove local or biological causality."
    elif hasattr(model, "coef_"):
        source = "linear_model_coefficient"
        coef = model.coef_
        if hasattr(coef, "tolist"):
            coef = coef.tolist()
        if coef and isinstance(coef[0], list):
            coef = coef[0]
        values = [abs(float(item)) for item in coef]
        summary = "The model exposes linear coefficients. These are coefficient magnitudes only and do not prove local or biological causality."
    else:
        warnings.append("Model-native feature importance or coefficients are not available for this estimator.")
        return [], "No model-native feature importance or coefficients are available for this model.", warnings

    pairs = []
    for index, feature in enumerate(FEATURE_COLUMNS):
        if index < len(values):
            pairs.append((feature, float(values[index])))
    pairs.sort(key=lambda item: abs(item[1]), reverse=True)
    important = [
        AdmetImportantFeature(
            feature=feature,
            value=round(value, 8),
            rank=rank,
            source=source,
            interpretation="Global model diagnostic only; not a causal or experimentally validated explanation.",
        )
        for rank, (feature, value) in enumerate(pairs[:10], start=1)
    ]
    return important, summary, warnings


def _external_validation_status(model_id: str, include_external_validation: bool) -> dict[str, Any]:
    if not include_external_validation:
        return {"status": "not_included", "reason": "External validation lookup was disabled for this explanation request."}
    latest = get_latest_external_validation_by_model(model_id)
    if not latest:
        return {"status": "not_available", "reason": "No external validation run is available for this model."}
    metrics = latest.get("metric_summary") or {}
    filtered_metrics = {key: value for key, value in metrics.items() if key not in {"observed_vs_predicted", "prediction_probabilities"}}
    warnings = latest.get("warnings") or []
    status = "available"
    if any("overfitting" in warning.lower() or "poorly calibrated" in warning.lower() for warning in warnings):
        status = "weak_or_concerning"
    return {
        "status": status,
        "run_id": latest.get("id"),
        "external_dataset_id": latest.get("external_dataset_id"),
        "valid_count": latest.get("valid_count"),
        "metric_summary": filtered_metrics,
        "calibration_summary": latest.get("calibration_summary") or {},
        "warnings": warnings,
        "created_at": latest.get("created_at"),
    }


def _metric_is_reasonable(task_type: str, metrics: dict[str, Any]) -> bool:
    if task_type == "binary_classification":
        for key in ("roc_auc", "f1", "balanced_accuracy", "accuracy"):
            value = metrics.get(key)
            if isinstance(value, (int, float)) and value >= 0.7:
                return True
        return False
    value = metrics.get("r2")
    return isinstance(value, (int, float)) and value >= 0.4


def _evidence_strength(
    model_summary: dict[str, Any],
    metrics: dict[str, Any],
    domain_status: str,
    uncertainty_level: str,
    external_validation: dict[str, Any],
) -> str:
    if model_summary.get("status") != "valid":
        return "not_available"
    train_count = 0
    training_run_id = model_summary.get("training_run_id")
    if training_run_id:
        init_db()
        with get_connection() as connection:
            row = connection.execute("SELECT train_count, test_count FROM admet_training_runs WHERE id = ?", (training_run_id,)).fetchone()
            if row:
                train_count = int(row["train_count"] or 0) + int(row["test_count"] or 0)
    external_status = external_validation.get("status")
    if external_status == "available" and domain_status != "outside_domain" and uncertainty_level != "high":
        return "externally_supported"
    if external_status == "weak_or_concerning":
        return "externally_weak"
    if domain_status == "outside_domain" or uncertainty_level == "high":
        return "weak_internal"
    if train_count >= 100 and _metric_is_reasonable(str(model_summary.get("task_type")), metrics):
        return "strong_internal_only"
    if train_count >= 20 and metrics:
        return "moderate_internal_only"
    return "uncertain"


def _save_explanation(
    report: AdmetPredictionExplanationResponse,
    training_run_id: int | None,
    report_files: dict[str, str | None] | None = None,
) -> int:
    init_db()
    with get_connection() as connection:
        cursor = connection.execute(
            """
            INSERT INTO admet_prediction_explanations (
                model_id, training_run_id, smiles, canonical_smiles,
                prediction_summary_json, explanation_summary_json, evidence_strength,
                domain_status, uncertainty_level, report_files_json, warnings_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                report.model_id,
                training_run_id,
                report.query_smiles,
                report.canonical_smiles,
                json.dumps(
                    {
                        "prediction_label": report.prediction_label,
                        "prediction_value": report.prediction_value,
                        "prediction_probability": report.prediction_probability,
                        "task_name": report.task_name,
                        "task_type": report.task_type,
                    }
                ),
                report.model_dump_json(),
                report.evidence_strength,
                report.domain_status,
                report.uncertainty_level,
                json.dumps(report_files or {}),
                json.dumps(report.warnings),
            ),
        )
        return int(cursor.lastrowid)


def explain_prediction(payload: AdmetPredictionExplainRequest) -> AdmetPredictionExplanationResponse:
    model_id = _resolve_model_id(payload.model_id)
    model_summary = _model_summary(model_id)
    try:
        mol = parse_smiles(payload.smiles)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Invalid SMILES structure: {exc}")
    canonical_smiles = Chem.MolToSmiles(mol, canonical=True)

    prediction = predict_trained_model(canonical_smiles, model_id)
    descriptor_values = _descriptor_features(canonical_smiles)
    folder = Path(model_summary["artifact_dir"])
    model_card = _read_json(folder / "model_card.json")
    training_summary = _read_json(folder / "training_summary.json")
    manifest = _read_json(folder / "model_manifest.json")
    metrics = training_summary.get("metrics") or manifest.get("metrics") or model_card.get("metrics") or {}

    warnings: list[str] = list(prediction.get("warnings") or [])
    descriptor_explanations, descriptor_warnings = _descriptor_explanations(model_id, descriptor_values)
    warnings.extend(w for w in descriptor_warnings if w not in warnings)
    important_features, contribution_summary, feature_warnings = _important_features(model_summary)
    warnings.extend(w for w in feature_warnings if w not in warnings)

    domain_status = prediction.get("domain_status") or "not_available"
    uncertainty_level = prediction.get("uncertainty_level") or "unknown"
    if payload.include_domain:
        try:
            domain = evaluate_domain_internal(model_id, canonical_smiles)
            domain_status = domain.get("domain_status", domain_status)
            uncertainty_level = domain.get("uncertainty_level", uncertainty_level)
            for warning in domain.get("warnings") or []:
                if warning not in warnings:
                    warnings.append(warning)
        except Exception as exc:
            warnings.append(f"Applicability domain explanation could not be refreshed: {exc}")
    else:
        warnings.append("Applicability domain was not included by request.")

    external_validation = _external_validation_status(model_id, payload.include_external_validation)
    if external_validation.get("status") == "not_available":
        warnings.append("No external validation summary is available for this model.")

    evidence_strength = _evidence_strength(model_summary, metrics, domain_status, uncertainty_level, external_validation)
    limitations = list(dict.fromkeys((prediction.get("limitations") or []) + (model_card.get("limitations") or []) + LIMITATIONS))

    response = AdmetPredictionExplanationResponse(
        model_id=model_id,
        model_name=prediction.get("model_name") or model_summary.get("model_name") or model_id,
        task_name=prediction.get("task_name") or model_summary.get("task_name"),
        task_type=prediction.get("task_type") or model_summary.get("task_type") or "unknown",
        query_smiles=payload.smiles,
        canonical_smiles=canonical_smiles,
        prediction_label=prediction.get("prediction_label"),
        prediction_value=prediction.get("prediction_value"),
        prediction_probability=prediction.get("prediction_score"),
        descriptor_values=descriptor_values,
        descriptor_explanations=descriptor_explanations,
        important_features=important_features,
        feature_contribution_summary=contribution_summary,
        domain_status=domain_status,
        uncertainty_level=uncertainty_level,
        external_validation_status=external_validation,
        evidence_strength=evidence_strength,
        model_card_summary=model_card,
        training_summary=training_summary,
        metrics=metrics,
        warnings=list(dict.fromkeys(warnings)),
        limitations=limitations,
        scientific_notice=SCIENTIFIC_NOTICE,
    )

    explanation_id = _save_explanation(response, model_summary.get("training_run_id"))
    if payload.project_id:
        _attach_explanation_to_project(payload.project_id, explanation_id, response)
    return response


def _pairs_table(data: dict[str, Any]) -> list[list[str]]:
    return [[str(key), json.dumps(value, default=str) if isinstance(value, (dict, list)) else str(value)] for key, value in data.items()]


def _build_pdf(report: AdmetPredictionExplanationResponse, path: Path) -> None:
    styles = getSampleStyleSheet()
    story = [
        Paragraph("DrugScreen360 ADMET Prediction Explanation Report", styles["Title"]),
        Paragraph(SCIENTIFIC_NOTICE, styles["Normal"]),
        Spacer(1, 12),
        Paragraph("Prediction Summary", styles["Heading2"]),
        Table(_pairs_table({
            "Model": f"{report.model_name} ({report.model_id})",
            "Task": f"{report.task_name or 'not available'} / {report.task_type}",
            "SMILES": report.canonical_smiles,
            "Prediction label": report.prediction_label or "not available",
            "Prediction value": report.prediction_value if report.prediction_value is not None else "not available",
            "Probability": report.prediction_probability if report.prediction_probability is not None else "not available",
            "Evidence strength": report.evidence_strength,
            "Domain status": report.domain_status,
            "Uncertainty": report.uncertainty_level,
        }), colWidths=[150, 360]),
        Spacer(1, 12),
        Paragraph("Important Features", styles["Heading2"]),
        Table(
            [["Rank", "Feature", "Value", "Source"]] + [[item.rank, item.feature, item.value, item.source] for item in report.important_features] or [["Status", "No model-native feature importance or coefficients available.", "", ""]],
            colWidths=[45, 145, 90, 220],
        ),
        Spacer(1, 12),
        Paragraph("Descriptor Context", styles["Heading2"]),
        Table(
            [["Feature", "Query", "Training min", "Training max", "Status"]]
            + [[item.feature, item.query_value, item.training_min, item.training_max, item.status] for item in report.descriptor_explanations[:12]],
            colWidths=[120, 80, 80, 80, 145],
        ),
        Spacer(1, 12),
        Paragraph("Warnings", styles["Heading2"]),
        Paragraph("; ".join(report.warnings) or "None", styles["Normal"]),
        Spacer(1, 12),
        Paragraph("Limitations", styles["Heading2"]),
        Paragraph("; ".join(report.limitations), styles["Normal"]),
    ]
    for element in story:
        if isinstance(element, Table):
            element.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]))
    SimpleDocTemplate(str(path), pagesize=letter).build(story)


def _build_docx(report: AdmetPredictionExplanationResponse, path: Path) -> None:
    document = Document()
    document.add_heading("DrugScreen360 ADMET Prediction Explanation Report", 0)
    document.add_paragraph(SCIENTIFIC_NOTICE)
    document.add_heading("Prediction Summary", level=1)
    for label, value in {
        "Model": f"{report.model_name} ({report.model_id})",
        "Task": f"{report.task_name or 'not available'} / {report.task_type}",
        "Canonical SMILES": report.canonical_smiles,
        "Prediction label": report.prediction_label or "not available",
        "Prediction value": report.prediction_value if report.prediction_value is not None else "not available",
        "Probability": report.prediction_probability if report.prediction_probability is not None else "not available",
        "Evidence strength": report.evidence_strength,
        "Domain status": report.domain_status,
        "Uncertainty": report.uncertainty_level,
    }.items():
        document.add_paragraph(f"{label}: {value}")
    document.add_heading("Important Features", level=1)
    if report.important_features:
        table = document.add_table(rows=1, cols=4)
        for idx, header in enumerate(["Rank", "Feature", "Value", "Source"]):
            table.rows[0].cells[idx].text = header
        for item in report.important_features:
            row = table.add_row().cells
            row[0].text = str(item.rank)
            row[1].text = item.feature
            row[2].text = str(item.value)
            row[3].text = item.source
    else:
        document.add_paragraph("No model-native feature importance or coefficients are available.")
    document.add_heading("Descriptor Context", level=1)
    for item in report.descriptor_explanations:
        document.add_paragraph(f"{item.feature}: {item.status}. {item.explanation}")
    document.add_heading("Warnings", level=1)
    for warning in report.warnings or ["None"]:
        document.add_paragraph(warning, style="List Bullet")
    document.add_heading("Limitations", level=1)
    for limitation in report.limitations:
        document.add_paragraph(limitation, style="List Bullet")
    document.save(path)


def create_explanation_report(payload: AdmetExplanationReportCreateRequest) -> AdmetExplanationReportCreateResponse:
    EXPLANATION_REPORT_DIR.mkdir(parents=True, exist_ok=True)
    explanation = explain_prediction(
        AdmetPredictionExplainRequest(
            model_id=payload.model_id,
            smiles=payload.smiles,
            include_domain=True,
            include_external_validation=True,
            project_id=payload.project_id,
        )
    )
    created = _now()
    stub = created.replace(":", "-").replace(".", "-")
    base = f"admet_explanation_{explanation.model_id}_{stub}"
    files: dict[str, str | None] = {"json": None, "pdf": None, "docx": None}
    formats = list(dict.fromkeys(payload.formats or ["json"]))
    if "json" in formats:
        path = EXPLANATION_REPORT_DIR / f"{base}.json"
        path.write_text(explanation.model_dump_json(indent=2), encoding="utf-8")
        files["json"] = path.name
    if "pdf" in formats:
        path = EXPLANATION_REPORT_DIR / f"{base}.pdf"
        _build_pdf(explanation, path)
        files["pdf"] = path.name
    if "docx" in formats:
        path = EXPLANATION_REPORT_DIR / f"{base}.docx"
        _build_docx(explanation, path)
        files["docx"] = path.name
    report_id = _save_explanation(explanation, _model_summary(explanation.model_id).get("training_run_id"), files)
    if payload.project_id:
        _attach_explanation_to_project(payload.project_id, report_id, explanation, files)
    return AdmetExplanationReportCreateResponse(
        report_id=report_id,
        model_id=explanation.model_id,
        created_at=created,
        available_formats=[fmt for fmt, filename in files.items() if filename],
        json_url=f"/api/admet-explain/reports/{report_id}/json" if files.get("json") else None,
        pdf_url=f"/api/admet-explain/reports/{report_id}/pdf" if files.get("pdf") else None,
        docx_url=f"/api/admet-explain/reports/{report_id}/docx" if files.get("docx") else None,
        warnings=explanation.warnings,
    )


def _attach_explanation_to_project(project_id: int, explanation_id: int, report: AdmetPredictionExplanationResponse, files: dict[str, str | None] | None = None) -> None:
    try:
        attach_project_item(
            project_id,
            ProjectAttachRequest(
                item_type="admet_prediction_explanation",
                item_id=str(explanation_id),
                item_title=f"ADMET Explanation: {report.model_name} on {report.canonical_smiles}",
                metadata={
                    "workflow_type": "admet_prediction_explainability",
                    "model_id": report.model_id,
                    "task_name": report.task_name,
                    "task_type": report.task_type,
                    "canonical_smiles": report.canonical_smiles,
                    "prediction_label": report.prediction_label,
                    "prediction_value": report.prediction_value,
                    "prediction_probability": report.prediction_probability,
                    "domain_status": report.domain_status,
                    "uncertainty_level": report.uncertainty_level,
                    "evidence_strength": report.evidence_strength,
                    "warnings": report.warnings,
                    "report_files": files or {},
                    "created_at": _now(),
                },
            ),
        )
    except Exception:
        pass


def list_explanation_reports() -> list[dict[str, Any]]:
    init_db()
    with get_connection() as connection:
        rows = connection.execute(
            "SELECT * FROM admet_prediction_explanations ORDER BY datetime(created_at) DESC, id DESC LIMIT 50"
        ).fetchall()
    items = []
    for row in rows:
        files = json.loads(row["report_files_json"]) if row["report_files_json"] else {}
        formats = [fmt for fmt, filename in files.items() if filename]
        items.append({
            "report_id": row["id"],
            "model_id": row["model_id"],
            "canonical_smiles": row["canonical_smiles"],
            "evidence_strength": row["evidence_strength"],
            "domain_status": row["domain_status"],
            "uncertainty_level": row["uncertainty_level"],
            "created_at": row["created_at"],
            "available_formats": formats,
            "json_url": f"/api/admet-explain/reports/{row['id']}/json" if files.get("json") else None,
            "pdf_url": f"/api/admet-explain/reports/{row['id']}/pdf" if files.get("pdf") else None,
            "docx_url": f"/api/admet-explain/reports/{row['id']}/docx" if files.get("docx") else None,
        })
    return items


def explanation_report_path(report_id: int, fmt: str) -> Path | None:
    init_db()
    with get_connection() as connection:
        row = connection.execute("SELECT report_files_json FROM admet_prediction_explanations WHERE id = ?", (report_id,)).fetchone()
    if not row:
        return None
    files = json.loads(row["report_files_json"]) if row["report_files_json"] else {}
    filename = files.get(fmt)
    if not filename:
        return None
    path = EXPLANATION_REPORT_DIR / filename
    return path if path.exists() else None


def explanation_summary_counts() -> dict[str, Any]:
    init_db()
    with get_connection() as connection:
        total = connection.execute("SELECT COUNT(*) FROM admet_prediction_explanations").fetchone()[0]
        evidence_rows = connection.execute(
            "SELECT evidence_strength, COUNT(*) AS c FROM admet_prediction_explanations GROUP BY evidence_strength"
        ).fetchall()
    return {
        "explanation_report_count": int(total or 0),
        "evidence_strength_distribution": {row["evidence_strength"]: row["c"] for row in evidence_rows},
    }
