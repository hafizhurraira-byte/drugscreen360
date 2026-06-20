import csv
import json
from io import BytesIO, StringIO
from typing import Any

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.database import get_connection, init_db
from app.models.project_report_models import ProjectReportPayload, ProjectReportStored
from app.services.model_registry import model_status_response


def _comparison_rows(payload: ProjectReportPayload) -> list[dict[str, Any]]:
    return payload.batch_screening_results.get("comparison_table") or []


def _top_candidates(payload: ProjectReportPayload, limit: int = 3) -> list[dict[str, Any]]:
    rows = _comparison_rows(payload)
    return sorted(
        rows,
        key=lambda row: (
            {"Higher priority": 0, "Review priority": 1, "Requires optimization": 2, "Treat cautiously": 3}.get(
                row.get("final_candidate_priority"), 4
            ),
            -(row.get("evidence_score") or 0),
            row.get("overall_admet_tox_concern_score") or 100,
        ),
    )[:limit]


def build_project_summary(payload: ProjectReportPayload) -> dict[str, Any]:
    top = _top_candidates(payload, 1)
    top_candidate = top[0].get("compound") or top[0].get("molecule_chembl_id") if top else None
    risks = []
    for row in _comparison_rows(payload):
        if row.get("concern_level") == "High":
            risks.append(f"{row.get('compound') or row.get('molecule_chembl_id')}: high ADMET/Tox concern")
        if row.get("evidence_level") in {"Weak", "Uncertain"}:
            risks.append(f"{row.get('compound') or row.get('molecule_chembl_id')}: weak/uncertain evidence")
    return {
        "workflow_type": payload.workflow_type,
        "selected_disease": payload.disease.disease_name if payload.disease else None,
        "selected_target": payload.disease_target.gene_symbol if payload.disease_target else None,
        "selected_chembl_target": payload.chembl_target.target_chembl_id if payload.chembl_target else None,
        "retrieved_candidate_count": payload.retrieved_candidate_count,
        "screened_candidate_count": payload.screened_candidate_count,
        "top_candidate": top_candidate,
        "main_risks": list(dict.fromkeys(risks))[:8],
    }


def save_project_report(payload: ProjectReportPayload, title: str | None = None) -> int:
    init_db()
    summary = build_project_summary(payload)
    report_title = title or payload.title
    with get_connection() as connection:
        cursor = connection.execute(
            """
            INSERT INTO project_reports (
                title, workflow_type, disease_name, disease_id, target_symbol,
                chembl_target_id, candidate_count, screened_count, top_candidate, report_payload_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                report_title,
                payload.workflow_type,
                payload.disease.disease_name if payload.disease else None,
                payload.disease.disease_id if payload.disease else None,
                payload.disease_target.gene_symbol if payload.disease_target else None,
                payload.chembl_target.target_chembl_id if payload.chembl_target else None,
                payload.retrieved_candidate_count,
                payload.screened_candidate_count,
                summary["top_candidate"],
                payload.model_dump_json(),
            ),
        )
        return int(cursor.lastrowid)


def get_project_report(report_id: int) -> ProjectReportStored | None:
    init_db()
    with get_connection() as connection:
        row = connection.execute("SELECT * FROM project_reports WHERE id = ?", (report_id,)).fetchone()
    if row is None:
        return None
    payload = ProjectReportPayload.model_validate_json(row["report_payload_json"])
    return ProjectReportStored(
        id=row["id"],
        title=row["title"],
        workflow_type=row["workflow_type"],
        disease_name=row["disease_name"],
        disease_id=row["disease_id"],
        target_symbol=row["target_symbol"],
        chembl_target_id=row["chembl_target_id"],
        candidate_count=row["candidate_count"],
        screened_count=row["screened_count"],
        top_candidate=row["top_candidate"],
        payload=payload,
        created_at=row["created_at"],
    )


def _pairs(mapping: dict[str, Any]) -> list[list[str]]:
    return [[str(key).replace("_", " ").title(), str(value if value is not None else "Not available")] for key, value in mapping.items()]


def _pdf_table(rows: list[list[str]], widths: list[float] | None = None) -> Table:
    table = Table(rows, colWidths=widths or [2.2 * inch, 4.3 * inch], repeatRows=1 if len(rows) > 1 else 0)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2f5")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#c8d1dc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.2),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _candidate_table_rows(payload: ProjectReportPayload) -> list[list[str]]:
    rows = [
        [
            "Rank",
            "Compound",
            "ChEMBL ID",
            "Activity",
            "Evidence",
            "MW",
            "LogP",
            "TPSA",
            "Risk",
            "ADMET/Tox",
            "Decision",
            "Priority",
        ]
    ]
    for index, row in enumerate(_comparison_rows(payload), start=1):
        rows.append(
            [
                str(index),
                str(row.get("compound") or "Unnamed"),
                str(row.get("molecule_chembl_id") or "NA"),
                f"{row.get('activity_type') or 'NA'} {row.get('activity_value') or ''} {row.get('activity_units') or ''}",
                f"{row.get('evidence_level') or 'NA'} ({row.get('evidence_score') or 'NA'})",
                str(row.get("molecular_weight") or "NA"),
                str(row.get("logp") or "NA"),
                str(row.get("tpsa") or "NA"),
                str(row.get("developability_risk") or "NA"),
                f"{row.get('concern_level') or 'NA'} ({row.get('overall_admet_tox_concern_score') or 'NA'})",
                str(row.get("decision") or "NA"),
                str(row.get("final_candidate_priority") or row.get("recommended_next_step") or "NA"),
            ]
        )
    return rows


def build_project_pdf(payload: ProjectReportPayload) -> bytes:
    buffer = BytesIO()
    styles = getSampleStyleSheet()
    story = []
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=0.45 * inch, leftMargin=0.45 * inch)
    summary = build_project_summary(payload)

    story.append(Paragraph(payload.title, styles["Title"]))
    story.append(Paragraph(payload.disclaimer, styles["BodyText"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Executive Summary", styles["Heading2"]))
    story.append(_pdf_table(_pairs(summary)))
    story.append(Spacer(1, 10))

    if payload.similarity:
        story.append(Paragraph("Reference Compound And Similarity Search", styles["Heading2"]))
        story.append(_pdf_table(_pairs(payload.similarity.model_dump())))
        story.append(
            Paragraph(
                "Chemical similarity does not prove shared efficacy, safety, mechanism, or regulatory acceptability.",
                styles["BodyText"],
            )
        )
        story.append(Spacer(1, 10))

    if payload.disease:
        story.append(Paragraph("Disease Context", styles["Heading2"]))
        story.append(_pdf_table(_pairs(payload.disease.model_dump())))
        if payload.disease_target:
            story.append(_pdf_table(_pairs(payload.disease_target.model_dump())))
        story.append(Paragraph("Open Targets association scores prioritize relevance only and do not prove safety or efficacy.", styles["BodyText"]))
        story.append(Spacer(1, 10))

    if payload.workflow_type != "similarity_to_candidate":
        story.append(Paragraph("Target Selection", styles["Heading2"]))
        story.append(_pdf_table(_pairs(payload.chembl_target.model_dump() if payload.chembl_target else {})))
        story.append(Spacer(1, 10))
    story.append(Paragraph("Candidate Retrieval Summary", styles["Heading2"]))
    retrieval_text = (
        "Source: similarity search. Canonical SMILES are required; duplicates are removed; analogs are ranked by similarity, "
        "identifier availability, data completeness, and RDKit Lipinski/Veber preview."
        if payload.workflow_type == "similarity_to_candidate"
        else "Source: ChEMBL. nM values are preferred; canonical SMILES are required; duplicates are removed; IC50/Ki/Kd are preferred where available."
    )
    story.append(Paragraph(retrieval_text, styles["BodyText"]))
    story.append(_pdf_table(_pairs({"retrieved_candidate_count": payload.retrieved_candidate_count, "screened_candidate_count": payload.screened_candidate_count})))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Candidate Comparison Table", styles["Heading2"]))
    story.append(_pdf_table(_candidate_table_rows(payload), widths=[0.35 * inch, 0.8 * inch, 0.75 * inch, 0.65 * inch, 0.65 * inch, 0.45 * inch, 0.45 * inch, 0.45 * inch, 0.6 * inch, 0.65 * inch, 0.75 * inch, 0.85 * inch]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Top Candidate Detail", styles["Heading2"]))
    for row in _top_candidates(payload, 3):
        story.append(Paragraph(str(row.get("compound") or row.get("molecule_chembl_id")), styles["Heading3"]))
        story.append(_pdf_table(_pairs({
            "activity": f"{row.get('activity_type')} {row.get('activity_value')} {row.get('activity_units')}",
            "evidence": f"{row.get('evidence_level')} ({row.get('evidence_score')})",
            "admet_tox": f"{row.get('concern_level')} ({row.get('overall_admet_tox_concern_score')})",
            "warnings": "; ".join(row.get("evidence_warnings") or []),
            "recommended_action": row.get("recommended_next_step"),
            "decision": row.get("decision"),
        })))
    story.append(Paragraph("Required Experimental Test Plan", styles["Heading2"]))
    story.append(Paragraph("Chemistry/developability, in vitro ADME, drug-drug interaction, safety/toxicity, genotoxicity, and nonclinical studies should be selected based on the candidate risk profile.", styles["BodyText"]))
    story.append(Paragraph("Prediction Model Status", styles["Heading2"]))
    status = model_status_response()
    story.append(_pdf_table(_pairs({
        "available_models": ", ".join(model.model_id for model in status["available_models"]),
        "unavailable_models": ", ".join(model.model_id for model in status["unavailable_models"]),
        "prediction_source_used": "Rule-based ADMET/Tox baseline unless real adapters are configured",
    })))
    story.append(Paragraph("Final Recommendation", styles["Heading2"]))
    story.append(Paragraph(f"Top candidate: {summary.get('top_candidate') or 'Not available'}. Review high-risk and weak-evidence candidates before advancing.", styles["BodyText"]))
    story.append(Paragraph("Limitations", styles["Heading2"]))
    for item in payload.limitations or []:
        story.append(Paragraph(f"- {item}", styles["BodyText"]))
    story.append(Paragraph("References/Data Sources", styles["Heading2"]))
    sources = ["PubChem", "ChEMBL", "RDKit", "Internal DrugScreen360 rule-based logic"]
    if payload.workflow_type == "disease_to_candidate":
        sources.append("Open Targets")
    story.append(Paragraph("; ".join(sources), styles["BodyText"]))
    doc.build(story)
    return buffer.getvalue()


def build_project_docx(payload: ProjectReportPayload) -> bytes:
    document = Document()
    summary = build_project_summary(payload)
    document.add_heading(payload.title, 0)
    document.add_paragraph(payload.disclaimer)
    document.add_heading("Executive Summary", level=1)
    for key, value in summary.items():
        document.add_paragraph(f"{key.replace('_', ' ').title()}: {value if value else 'Not available'}")
    if payload.similarity:
        document.add_heading("Reference Compound And Similarity Search", level=1)
        for key, value in payload.similarity.model_dump().items():
            document.add_paragraph(f"{key.replace('_', ' ').title()}: {value or 'Not available'}")
        document.add_paragraph("Chemical similarity does not prove shared efficacy, safety, mechanism, or regulatory acceptability.")
    if payload.disease:
        document.add_heading("Disease Context", level=1)
        for key, value in payload.disease.model_dump().items():
            document.add_paragraph(f"{key.replace('_', ' ').title()}: {value or 'Not available'}")
        if payload.disease_target:
            for key, value in payload.disease_target.model_dump().items():
                document.add_paragraph(f"{key.replace('_', ' ').title()}: {value or 'Not available'}")
    if payload.workflow_type != "similarity_to_candidate":
        document.add_heading("Target Selection", level=1)
        for key, value in (payload.chembl_target.model_dump() if payload.chembl_target else {}).items():
            document.add_paragraph(f"{key.replace('_', ' ').title()}: {value or 'Not available'}")
    document.add_heading("Candidate Retrieval Summary", level=1)
    if payload.workflow_type == "similarity_to_candidate":
        document.add_paragraph("Source: similarity search. Canonical SMILES are required; duplicates are removed; analogs are ranked by similarity, identifier availability, data completeness, and RDKit Lipinski/Veber preview.")
    else:
        document.add_paragraph("Source: ChEMBL. nM values are preferred; canonical SMILES are required; duplicates are removed; IC50/Ki/Kd are preferred where available.")
    document.add_heading("Candidate Comparison Table", level=1)
    for row in _comparison_rows(payload):
        document.add_paragraph(
            f"{row.get('compound') or row.get('molecule_chembl_id')}: {row.get('activity_type')} {row.get('activity_value')} {row.get('activity_units')}; "
            f"Evidence {row.get('evidence_level')} ({row.get('evidence_score')}); ADMET/Tox {row.get('concern_level')} "
            f"({row.get('overall_admet_tox_concern_score')}); Decision {row.get('decision')}; Priority {row.get('final_candidate_priority')}",
            style="List Bullet",
        )
    document.add_heading("Top Candidate Detail", level=1)
    for row in _top_candidates(payload, 3):
        document.add_heading(str(row.get("compound") or row.get("molecule_chembl_id")), level=2)
        document.add_paragraph(f"Recommended action: {row.get('recommended_next_step') or 'Review with expert team.'}")
    document.add_heading("Required Experimental Test Plan", level=1)
    for item in ["Chemistry/developability", "In vitro ADME", "Drug-drug interaction", "Safety/toxicity", "Genotoxicity", "Nonclinical studies if moving forward"]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("Prediction Model Status", level=1)
    status = model_status_response()
    document.add_paragraph(f"Available models: {', '.join(model.model_id for model in status['available_models'])}")
    document.add_paragraph(f"Unavailable models: {', '.join(model.model_id for model in status['unavailable_models'])}")
    document.add_paragraph("Prediction source used: Rule-based ADMET/Tox baseline unless real adapters are configured.")
    document.add_heading("Final Recommendation", level=1)
    document.add_paragraph(f"Top candidate: {summary.get('top_candidate') or 'Not available'}. Review high-risk and weak-evidence candidates before advancing.")
    document.add_heading("Limitations", level=1)
    for item in payload.limitations or []:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("References/Data Sources", level=1)
    document.add_paragraph("PubChem; ChEMBL; Open Targets when disease workflow is used; RDKit; Internal DrugScreen360 rule-based logic")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_project_csv(payload: ProjectReportPayload) -> str:
    headers = [
        "compound", "pubchem_cid", "molecule_chembl_id", "similarity_score", "source", "target_name", "activity_type", "activity_value", "activity_units",
        "evidence_level", "evidence_score", "molecular_weight", "logp", "tpsa", "drug_likeness_status",
        "overall_admet_tox_concern_score", "concern_level", "decision", "recommended_next_step", "analog_priority_score",
    ]
    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=headers, extrasaction="ignore")
    writer.writeheader()
    for row in _comparison_rows(payload):
        writer.writerow(row)
    return output.getvalue()
