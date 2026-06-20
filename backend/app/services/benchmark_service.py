import csv
import json
from collections import Counter
from datetime import datetime, timezone
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import HTTPException
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.constants import DISCLAIMER
from app.database import get_connection, init_db
from app.models.benchmark_models import (
    BenchmarkCompound,
    BenchmarkResultItem,
    BenchmarkRunResponse,
    BenchmarkRunStored,
    BenchmarkSummary,
)
from app.models.schemas import CompoundIdentity
from app.services.admet_toxicity_engine import evaluate_admet_toxicity
from app.services.descriptors import calculate_descriptors, parse_smiles
from app.services.pubchem import resolve_compound
from app.services.rules import build_decision, evaluate_rules, plan_experimental_tests
from app.services.model_registry import model_status_response

BENCHMARK_DISCLAIMER = (
    "This benchmark checks internal rule behavior only. It does not validate clinical safety, efficacy, "
    "regulatory approval, or market readiness."
)
DATA_PATH = Path(__file__).resolve().parents[1] / "benchmark_data" / "benchmark_compounds.json"


def load_benchmark_groups() -> dict[str, list[BenchmarkCompound]]:
    raw = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    groups: dict[str, list[BenchmarkCompound]] = {}
    for group, items in raw.items():
        groups[group] = [BenchmarkCompound.model_validate({**item, "group": group}) for item in items]
    return groups


def all_benchmark_compounds() -> list[BenchmarkCompound]:
    return [item for items in load_benchmark_groups().values() for item in items]


def _direct_identity(item: BenchmarkCompound) -> CompoundIdentity:
    return CompoundIdentity(
        compound_name=item.name,
        pubchem_cid=None,
        canonical_smiles=item.query,
        isomeric_smiles=item.query,
        molecular_formula=None,
        molecular_weight=None,
        iupac_name=None,
        synonyms=[],
        pubchem_source_link=None,
    )


def _screen_item(item: BenchmarkCompound) -> dict[str, Any]:
    if item.input_type == "smiles":
        parse_smiles(item.query)
        identity = _direct_identity(item)
    else:
        identity = resolve_compound(item.query, item.input_type)
    smiles = identity.canonical_smiles or identity.isomeric_smiles
    if not smiles:
        raise HTTPException(status_code=422, detail="No usable SMILES returned for benchmark compound.")
    descriptors = calculate_descriptors(smiles)
    rules = evaluate_rules(descriptors)
    admet_tox = evaluate_admet_toxicity(smiles, descriptors)
    tests = plan_experimental_tests(descriptors, rules)
    decision = build_decision(rules, tests)
    return {
        "identity": identity,
        "descriptors": descriptors,
        "rules": rules,
        "admet_tox": admet_tox,
        "decision": decision,
    }


def _evaluate_expectation(item: BenchmarkCompound, output: dict[str, Any] | None, error: Exception | None) -> tuple[str, str, str]:
    expected = item.expected_warning_category
    if expected == "clean_validation_error":
        if error is not None and "Invalid SMILES" in str(error):
            return "PASS", "Invalid SMILES returned a clean validation error.", "Keep validation messaging clear."
        return "FAIL", "Invalid SMILES did not return the expected clean validation behavior.", "Inspect input validation path."

    if error is not None:
        return "FAIL", f"Benchmark item failed unexpectedly: {error}", "Check external lookup/cache availability or parser behavior."

    assert output is not None
    rules = output["rules"]
    admet_tox = output["admet_tox"]
    score = admet_tox.overall.overall_admet_tox_concern_score
    concern = admet_tox.overall.concern_level
    alerts = admet_tox.structural_alerts.structural_alerts

    if expected in {"low_concern", "acceptable"}:
        if rules.basic_drug_likeness_status != "Poor" and concern in {"Low", "Medium"}:
            return "PASS", "Rule output is broadly consistent with expected acceptable/low-concern behavior.", "No immediate rule change needed."
        if score < 70:
            return "REVIEW", "Output is valid but stronger than expected for this reference item.", "Review thresholds and descriptor reasons."
        return "FAIL", "Output appears much stronger than expected for this benchmark item.", "Inspect drug-likeness and ADMET/Tox thresholds."

    if expected == "high_developability_admet":
        if rules.developability_risk == "High" or concern == "High" or score >= 50:
            return "PASS", "Rule output raised higher developability or ADMET/Tox concern as expected.", "Review details but behavior is consistent."
        if rules.developability_risk == "Medium" or concern == "Medium" or score >= 35:
            return "REVIEW", "Output raised some caution but not a strong warning.", "Consider whether thresholds should be more sensitive."
        return "FAIL", "Expected stronger warnings were not observed.", "Inspect MW/LogP/TPSA and concern scoring rules."

    if expected == "structural_alert":
        if alerts or admet_tox.structural_alerts.structural_alert_risk in {"Medium", "High"}:
            return "PASS", "Structural alert behavior is present as expected.", "No immediate rule change needed."
        return "REVIEW", "Compound screened successfully but expected structural alert was not triggered.", "Review SMARTS alert coverage."

    return "REVIEW", "Expected category is not mapped to a strict benchmark rule.", "Add a benchmark evaluator rule if this case becomes important."


def evaluate_benchmark_item(item: BenchmarkCompound) -> BenchmarkResultItem:
    output = None
    error = None
    try:
        output = _screen_item(item)
    except Exception as exc:
        error = exc
    status, reason, recommendation = _evaluate_expectation(item, output, error)

    if output is None:
        return BenchmarkResultItem(
            benchmark_id=item.id,
            compound=item.name,
            group=item.group or "unknown",
            input_type=item.input_type,
            query=item.query,
            expected_behavior=item.expected_general_behavior,
            expected_warning_category=item.expected_warning_category,
            status=status,
            reason=reason,
            clean_error=str(error) if error else None,
            recommendation=recommendation,
        )

    descriptors = output["descriptors"]
    rules = output["rules"]
    admet_tox = output["admet_tox"]
    identity = output["identity"]
    return BenchmarkResultItem(
        benchmark_id=item.id,
        compound=item.name,
        group=item.group or "unknown",
        input_type=item.input_type,
        query=item.query,
        expected_behavior=item.expected_general_behavior,
        expected_warning_category=item.expected_warning_category,
        status=status,
        reason=reason,
        actual_decision=output["decision"]["decision"],
        drug_likeness=rules.basic_drug_likeness_status,
        developability_risk=rules.developability_risk,
        admet_tox_concern_score=admet_tox.overall.overall_admet_tox_concern_score,
        admet_tox_concern_level=admet_tox.overall.concern_level,
        structural_alert_risk=admet_tox.structural_alerts.structural_alert_risk,
        structural_alerts=admet_tox.structural_alerts.structural_alerts,
        descriptor_summary=descriptors.model_dump(),
        recommendation=recommendation,
        cache_metadata=identity.cache_metadata.model_dump() if identity.cache_metadata else None,
    )


def _select_items(selected_ids: list[str], group_name: str | None, max_items: int | None) -> list[BenchmarkCompound]:
    groups = load_benchmark_groups()
    if group_name:
        selected = groups.get(group_name)
        if selected is None:
            raise HTTPException(status_code=404, detail=f"Benchmark group '{group_name}' not found.")
    elif selected_ids:
        by_id = {item.id: item for item in all_benchmark_compounds()}
        selected = [by_id[item_id] for item_id in selected_ids if item_id in by_id]
    else:
        selected = all_benchmark_compounds()
    if max_items:
        selected = selected[:max_items]
    if not selected:
        raise HTTPException(status_code=422, detail="No benchmark compounds selected.")
    return selected


def _summary(results: list[BenchmarkResultItem]) -> BenchmarkSummary:
    warning_counter: Counter[str] = Counter()
    for item in results:
        if item.developability_risk in {"Medium", "High"}:
            warning_counter[item.developability_risk + " developability risk"] += 1
        if item.admet_tox_concern_level in {"Medium", "High"}:
            warning_counter[item.admet_tox_concern_level + " ADMET/Tox concern"] += 1
        if item.structural_alerts:
            warning_counter["Structural alerts"] += 1
        if item.clean_error:
            warning_counter["Clean validation errors"] += 1
    return BenchmarkSummary(
        total_tested=len(results),
        passed=sum(1 for item in results if item.status == "PASS"),
        review=sum(1 for item in results if item.status == "REVIEW"),
        failed=sum(1 for item in results if item.status == "FAIL"),
        warning_count=sum(warning_counter.values()),
        most_common_warnings=[name for name, _ in warning_counter.most_common(6)],
        compounds_needing_review=[item.compound for item in results if item.status in {"REVIEW", "FAIL"}],
    )


def run_benchmark(selected_ids: list[str], group_name: str | None, max_items: int | None) -> BenchmarkRunResponse:
    selected = _select_items(selected_ids, group_name, max_items)
    results = [evaluate_benchmark_item(item) for item in selected]
    summary = _summary(results)
    model_status = model_status_response()
    external_info = next(
        (model for model in model_status["available_models"] + model_status["unavailable_models"] if model.model_id == "external_admet_provider_v1"),
        None,
    )
    response = BenchmarkRunResponse(
        selected_group=group_name,
        summary=summary,
        individual_results=results,
        mismatches=[item for item in results if item.status in {"REVIEW", "FAIL"}],
        limitations=[
            BENCHMARK_DISCLAIMER,
            DISCLAIMER,
            "Benchmark expectations are broad internal checks, not clinical truth labels.",
            "Named compounds may depend on PubChem/cache availability.",
        ],
        model_status_summary={
            "used_models": ["rule_based_admet_v1"],
            "only_rule_based_output_used": True,
            "unavailable_model_count": len(model_status["unavailable_models"]),
            "external_provider_status": external_info.status if external_info else "not_registered",
            "external_model_available": bool(external_info and external_info.status == "available"),
            "mock_provider_used": bool(external_info and external_info.status == "mock"),
            "external_model_warning": external_info.warning if external_info else "External ADMET provider adapter is not registered.",
            "message": "Only rule-based ADMET/Tox screening is active unless a real adapter is configured.",
        },
        created_at=datetime.now(timezone.utc).isoformat(),
    )
    response.benchmark_run_id = save_benchmark_run(response)
    return response


def save_benchmark_run(response: BenchmarkRunResponse) -> int:
    init_db()
    with get_connection() as connection:
        cursor = connection.execute(
            """
            INSERT INTO benchmark_runs (
                title, selected_group, total_tested, passed, review, failed, result_payload_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                response.title,
                response.selected_group,
                response.summary.total_tested,
                response.summary.passed,
                response.summary.review,
                response.summary.failed,
                response.model_dump_json(),
            ),
        )
        return int(cursor.lastrowid)


def list_benchmark_runs() -> list[BenchmarkRunStored]:
    init_db()
    with get_connection() as connection:
        rows = connection.execute("SELECT * FROM benchmark_runs ORDER BY datetime(created_at) DESC, id DESC LIMIT 50").fetchall()
    return [_stored_from_row(row) for row in rows]


def get_benchmark_run(run_id: int) -> BenchmarkRunStored | None:
    init_db()
    with get_connection() as connection:
        row = connection.execute("SELECT * FROM benchmark_runs WHERE id = ?", (run_id,)).fetchone()
    return _stored_from_row(row) if row else None


def _stored_from_row(row) -> BenchmarkRunStored:
    return BenchmarkRunStored(
        id=row["id"],
        title=row["title"],
        selected_group=row["selected_group"],
        total_tested=row["total_tested"],
        passed=row["passed"],
        review=row["review"],
        failed=row["failed"],
        payload=BenchmarkRunResponse.model_validate_json(row["result_payload_json"]),
        created_at=row["created_at"],
    )


def benchmark_csv(response: BenchmarkRunResponse) -> str:
    headers = [
        "compound",
        "group",
        "expected_warning_category",
        "actual_decision",
        "drug_likeness",
        "developability_risk",
        "admet_tox_concern_score",
        "admet_tox_concern_level",
        "status",
        "reason",
        "recommendation",
    ]
    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=headers, extrasaction="ignore")
    writer.writeheader()
    for item in response.individual_results:
        writer.writerow(item.model_dump())
    return output.getvalue()


def _pdf_table(rows: list[list[str]], widths: list[float] | None = None) -> Table:
    table = Table(rows, colWidths=widths or [1.5 * inch, 5.0 * inch], repeatRows=1 if len(rows) > 1 else 0)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2f5")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#c8d1dc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def benchmark_pdf(response: BenchmarkRunResponse) -> bytes:
    buffer = BytesIO()
    styles = getSampleStyleSheet()
    story = []
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=0.45 * inch, leftMargin=0.45 * inch)
    story.append(Paragraph("DrugScreen360 Validation & Benchmarking Report", styles["Title"]))
    story.append(Paragraph(BENCHMARK_DISCLAIMER, styles["BodyText"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Executive Summary", styles["Heading2"]))
    story.append(_pdf_table([[k.replace("_", " ").title(), str(v)] for k, v in response.summary.model_dump().items()]))
    story.append(Paragraph("Benchmark Dataset Description", styles["Heading2"]))
    story.append(Paragraph("Local benchmark groups include common reference drugs, warning compounds, and chemistry stress tests.", styles["BodyText"]))
    story.append(Paragraph("Method", styles["Heading2"]))
    story.append(Paragraph("Each compound is screened with the existing descriptor, drug-likeness, and ADMET/Tox rule services, then compared with broad expected behavior.", styles["BodyText"]))
    story.append(Paragraph("Prediction Model Status", styles["Heading2"]))
    story.append(_pdf_table([[k.replace("_", " ").title(), str(v)] for k, v in response.model_status_summary.items()]))
    story.append(Paragraph("Benchmark Results Table", styles["Heading2"]))
    rows = [["Compound", "Group", "Expected", "Decision", "Concern", "Status"]]
    for item in response.individual_results:
        rows.append([
            item.compound,
            item.group,
            item.expected_warning_category,
            item.actual_decision or item.clean_error or "NA",
            str(item.admet_tox_concern_level or "NA"),
            item.status,
        ])
    story.append(_pdf_table(rows, widths=[1.1 * inch, 1.15 * inch, 1.25 * inch, 1.25 * inch, 0.9 * inch, 0.65 * inch]))
    story.append(Paragraph("Individual Compound Results", styles["Heading2"]))
    for item in response.individual_results:
        story.append(Paragraph(item.compound, styles["Heading3"]))
        story.append(Paragraph(f"{item.status}: {item.reason}", styles["BodyText"]))
    story.append(Paragraph("Mismatches / Needs Review", styles["Heading2"]))
    story.append(Paragraph("; ".join(item.compound for item in response.mismatches) or "None", styles["BodyText"]))
    story.append(Paragraph("Limitations", styles["Heading2"]))
    for item in response.limitations:
        story.append(Paragraph(f"- {item}", styles["BodyText"]))
    story.append(Paragraph("Recommendations for Rule Improvement", styles["Heading2"]))
    story.append(Paragraph("Review REVIEW/FAIL cases, inspect descriptor thresholds, and expand structural alert SMARTS only with documented rationale.", styles["BodyText"]))
    doc.build(story)
    return buffer.getvalue()


def benchmark_docx(response: BenchmarkRunResponse) -> bytes:
    document = Document()
    document.add_heading("DrugScreen360 Validation & Benchmarking Report", 0)
    document.add_paragraph(BENCHMARK_DISCLAIMER)
    document.add_heading("Executive Summary", level=1)
    for key, value in response.summary.model_dump().items():
        document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
    document.add_heading("Benchmark Dataset Description", level=1)
    document.add_paragraph("Local benchmark groups include common reference drugs, warning compounds, and chemistry stress tests.")
    document.add_heading("Method", level=1)
    document.add_paragraph("Each compound is screened with existing transparent rules and compared with broad expected behavior.")
    document.add_heading("Prediction Model Status", level=1)
    for key, value in response.model_status_summary.items():
        document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
    document.add_heading("Benchmark Results Table", level=1)
    for item in response.individual_results:
        document.add_paragraph(f"{item.compound}: {item.status} - {item.reason}", style="List Bullet")
    document.add_heading("Mismatches / Needs Review", level=1)
    for item in response.mismatches:
        document.add_paragraph(f"{item.compound}: {item.recommendation}", style="List Bullet")
    document.add_heading("Limitations", level=1)
    for item in response.limitations:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("Recommendations for Rule Improvement", level=1)
    document.add_paragraph("Review REVIEW/FAIL cases, inspect descriptor thresholds, and expand structural alert SMARTS only with documented rationale.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
