import csv
import json
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import HTTPException
from rdkit import Chem
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.constants import DISCLAIMER
from app.database import get_connection, init_db
from app.models.batch_library_models import BatchLibraryResultRow, BatchLibraryScreenResponse, BatchLibraryRunStored, BatchParseResponse, ParsedCompound
from app.services.admet_predictor_service import predict_admet
from app.services.admet_toxicity_engine import evaluate_admet_toxicity
from app.services.descriptors import calculate_descriptors, parse_smiles
from app.services.model_registry import model_status_response
from app.services.rules import build_decision, evaluate_rules, plan_experimental_tests

MAX_FILE_SIZE = 5 * 1024 * 1024
MAX_PARSE_COMPOUNDS = 500
ALLOWED_EXTENSIONS = {".csv", ".txt", ".smi", ".sdf", ".mol"}
LIMITATION = (
    "Uploaded compounds are screened using computational descriptors and rule-based/model-adapter outputs only. "
    "Results do not prove safety, efficacy, clinical success, regulatory approval, or market readiness."
)


def _canonicalize(smiles: str) -> str:
    mol = parse_smiles(smiles.strip())
    return Chem.MolToSmiles(mol, canonical=True)


def _descriptor_dict(smiles: str) -> dict[str, Any]:
    return calculate_descriptors(smiles).model_dump()


def _normalized(row: dict[str, Any], *names: str) -> str | None:
    lowered = {str(key).strip().lower(): value for key, value in row.items()}
    for name in names:
        if name in lowered and lowered[name] not in {None, ""}:
            return str(lowered[name]).strip()
    return None


def _parse_csv(content: bytes) -> list[dict[str, Any]]:
    text = content.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(StringIO(text))
    if not reader.fieldnames:
        raise HTTPException(status_code=422, detail="CSV file is empty or missing a header row.")
    rows = []
    for index, row in enumerate(reader, start=2):
        rows.append(
            {
                "row_number": index,
                "smiles": _normalized(row, "smiles", "canonical_smiles", "molecule_smiles"),
                "name": _normalized(row, "name", "compound_name", "molecule_name"),
                "compound_id": _normalized(row, "compound_id", "id"),
                "source": _normalized(row, "source"),
                "notes": _normalized(row, "notes"),
            }
        )
    return rows


def _parse_smi(content: bytes) -> list[dict[str, Any]]:
    text = content.decode("utf-8-sig", errors="replace")
    rows = []
    for index, line in enumerate(text.splitlines(), start=1):
        clean = line.strip()
        if not clean or clean.startswith("#"):
            continue
        parts = clean.split()
        rows.append({"row_number": index, "smiles": parts[0], "name": " ".join(parts[1:]) or None})
    return rows


def _parse_sdf(content: bytes) -> list[dict[str, Any]]:
    supplier = Chem.ForwardSDMolSupplier(BytesIO(content), sanitize=True, removeHs=False)
    rows = []
    for index, mol in enumerate(supplier, start=1):
        if mol is None:
            rows.append({"row_number": index, "smiles": None, "name": None, "error_reason": "RDKit could not parse SDF molecule."})
            continue
        props = mol.GetPropsAsDict()
        rows.append(
            {
                "row_number": index,
                "smiles": Chem.MolToSmiles(mol, canonical=True),
                "name": mol.GetProp("_Name") if mol.HasProp("_Name") else None,
                "compound_id": str(props.get("compound_id") or props.get("id") or "") or None,
                "source": str(props.get("source") or "") or None,
                "notes": json.dumps(props) if props else None,
            }
        )
    return rows


def _parse_mol(content: bytes) -> list[dict[str, Any]]:
    mol = Chem.MolFromMolBlock(content.decode("utf-8", errors="replace"), sanitize=True, removeHs=False)
    if mol is None:
        return [{"row_number": 1, "smiles": None, "error_reason": "RDKit could not parse MOL file."}]
    return [{"row_number": 1, "smiles": Chem.MolToSmiles(mol, canonical=True), "name": mol.GetProp("_Name") if mol.HasProp("_Name") else None}]


def parse_library_file(file_name: str, content: bytes) -> BatchParseResponse:
    if not content:
        raise HTTPException(status_code=422, detail="Uploaded file is empty.")
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File is too large. Maximum size is 5 MB.")
    extension = Path(file_name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=415, detail="Unsupported file type. Use CSV, TXT, SMI, SDF, or MOL.")

    if extension == ".csv":
        raw_rows = _parse_csv(content)
    elif extension in {".txt", ".smi"}:
        raw_rows = _parse_smi(content)
    elif extension == ".sdf":
        raw_rows = _parse_sdf(content)
    else:
        raw_rows = _parse_mol(content)

    if not raw_rows:
        raise HTTPException(status_code=422, detail="No compounds found in uploaded file.")
    if len(raw_rows) > MAX_PARSE_COMPOUNDS:
        raise HTTPException(status_code=422, detail=f"Too many compounds. Maximum parse limit is {MAX_PARSE_COMPOUNDS}.")

    seen: set[str] = set()
    parsed = []
    warnings = []
    duplicate_count = 0
    for row in raw_rows:
        original = (row.get("smiles") or "").strip()
        error = row.get("error_reason")
        canonical = None
        descriptors = None
        valid = False
        duplicate = False
        if not error:
            if not original:
                error = "Missing SMILES."
            else:
                try:
                    canonical = _canonicalize(original)
                    descriptors = _descriptor_dict(canonical)
                    valid = True
                    if canonical in seen:
                        duplicate = True
                        duplicate_count += 1
                    seen.add(canonical)
                except Exception as exc:
                    error = str(exc)
        parsed.append(
            ParsedCompound(
                row_number=int(row.get("row_number") or len(parsed) + 1),
                compound_name=row.get("name"),
                compound_id=row.get("compound_id"),
                original_smiles=original or None,
                canonical_smiles=canonical,
                valid=valid,
                duplicate=duplicate,
                error_reason=error,
                descriptors=descriptors,
                source=row.get("source"),
                notes=row.get("notes"),
            )
        )
    batch_id = save_parsed_batch(file_name, extension.lstrip("."), parsed)
    valid_count = sum(1 for item in parsed if item.valid)
    invalid_count = sum(1 for item in parsed if not item.valid)
    if duplicate_count:
        warnings.append(f"{duplicate_count} duplicate canonical SMILES detected.")
    return BatchParseResponse(
        batch_id=batch_id,
        file_name=file_name,
        file_type=extension.lstrip("."),
        total_rows=len(raw_rows),
        valid_compounds=valid_count,
        invalid_compounds=invalid_count,
        duplicates_detected=duplicate_count,
        parsed_compounds=parsed,
        warnings=warnings,
        limitations=[LIMITATION, "Evidence quality is not evaluated unless uploaded compounds are linked to target bioactivity records."],
    )


def save_parsed_batch(file_name: str, file_type: str, parsed: list[ParsedCompound]) -> int:
    init_db()
    with get_connection() as connection:
        cursor = connection.execute(
            """
            INSERT INTO uploaded_batches (file_name, file_type, total_rows, valid_count, invalid_count, duplicate_count)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                file_name,
                file_type,
                len(parsed),
                sum(1 for item in parsed if item.valid),
                sum(1 for item in parsed if not item.valid),
                sum(1 for item in parsed if item.duplicate),
            ),
        )
        batch_id = int(cursor.lastrowid)
        for item in parsed:
            connection.execute(
                """
                INSERT INTO uploaded_batch_compounds (
                    batch_id, row_number, compound_name, compound_id, original_smiles,
                    canonical_smiles, valid, error_reason, descriptors_json, source, notes
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    batch_id,
                    item.row_number,
                    item.compound_name,
                    item.compound_id,
                    item.original_smiles,
                    item.canonical_smiles,
                    1 if item.valid else 0,
                    item.error_reason,
                    json.dumps(item.descriptors) if item.descriptors else None,
                    item.source,
                    item.notes,
                ),
            )
    return batch_id


def get_batch_compounds(batch_id: int) -> list[ParsedCompound]:
    init_db()
    with get_connection() as connection:
        rows = connection.execute("SELECT * FROM uploaded_batch_compounds WHERE batch_id = ? ORDER BY row_number, id", (batch_id,)).fetchall()
    return [
        ParsedCompound(
            row_number=row["row_number"],
            compound_name=row["compound_name"],
            compound_id=row["compound_id"],
            original_smiles=row["original_smiles"],
            canonical_smiles=row["canonical_smiles"],
            valid=bool(row["valid"]),
            error_reason=row["error_reason"],
            descriptors=json.loads(row["descriptors_json"]) if row["descriptors_json"] else None,
            source=row["source"],
            notes=row["notes"],
        )
        for row in rows
    ]


def _priority_score(rules, admet, model_status: str) -> tuple[float, str]:
    score = 100
    score -= {"Low": 0, "Medium": 18, "High": 35}.get(rules.developability_risk, 20)
    score -= min(admet.overall.overall_admet_tox_concern_score * 0.45, 40)
    if not rules.lipinski_rule_of_5["passed"]:
        score -= 12
    if not rules.veber_rule["passed"]:
        score -= 10
    if model_status != "available":
        score -= 5
    score = round(max(score, 0), 2)
    label = "High" if score >= 75 else "Medium" if score >= 55 else "Review" if score >= 35 else "Low"
    return score, label


def screen_batch_library(batch_id: int | None, compounds: list[ParsedCompound], max_compounds: int, run_model_predictions: bool) -> BatchLibraryScreenResponse:
    selected = get_batch_compounds(batch_id) if batch_id else compounds
    valid = [item for item in selected if item.valid and item.canonical_smiles]
    if not valid:
        raise HTTPException(status_code=422, detail="No valid compounds available for screening.")
    if len(valid) > max_compounds:
        raise HTTPException(status_code=422, detail=f"Too many valid compounds selected. Maximum screening limit is {max_compounds}.")

    results = []
    failed_count = 0
    warnings = []
    for item in valid:
        try:
            descriptors = calculate_descriptors(item.canonical_smiles)
            rules = evaluate_rules(descriptors)
            admet = evaluate_admet_toxicity(item.canonical_smiles, descriptors)
            tests = plan_experimental_tests(descriptors, rules)
            decision = build_decision(rules, tests)
            model_predictions = (
                predict_admet(item.canonical_smiles, ["rule_based_admet_v1", "external_admet_provider_v1", "local_admet_model", "trained_local_admet_model"], True)
                if run_model_predictions
                else None
            )
            rule_model = model_predictions.model_outputs[0] if model_predictions and model_predictions.model_outputs else None
            model_summary = model_predictions.model_status_summary if model_predictions else {}
            model_status = rule_model.model_status if rule_model else "not_run"
            score, label = _priority_score(rules, admet, model_status)
            results.append(
                BatchLibraryResultRow(
                    row_number=item.row_number,
                    compound_name=item.compound_name,
                    compound_id=item.compound_id,
                    canonical_smiles=item.canonical_smiles,
                    molecular_weight=descriptors.molecular_weight,
                    logp=descriptors.logp,
                    tpsa=descriptors.tpsa,
                    lipinski_pass=bool(rules.lipinski_rule_of_5["passed"]),
                    veber_pass=bool(rules.veber_rule["passed"]),
                    drug_likeness_status=rules.basic_drug_likeness_status,
                    developability_risk=rules.developability_risk,
                    absorption_risk=admet.absorption.absorption_risk,
                    solubility_risk=admet.solubility.solubility_risk,
                    structural_alert_risk=admet.structural_alerts.structural_alert_risk,
                    overall_admet_tox_concern_score=admet.overall.overall_admet_tox_concern_score,
                    concern_level=admet.overall.concern_level,
                    decision=decision["decision"],
                    admet_prediction_source=rule_model.prediction_source if rule_model else "Not run",
                    model_status=model_status,
                    model_confidence=rule_model.confidence if rule_model else "Not run",
                    model_warnings=model_predictions.warnings if model_predictions else ["Model predictions were not run."],
                    rule_based_used=bool(model_summary.get("rule_based_used", run_model_predictions)),
                    external_model_used=bool(model_summary.get("external_model_used", False)),
                    external_model_available=bool(model_summary.get("external_model_available", False)),
                    external_model_warning=model_summary.get("external_model_warning"),
                    trained_model_used=bool(model_summary.get("trained_model_used", False)),
                    trained_model_available=bool(model_summary.get("trained_model_available", False)),
                    trained_model_warning=model_summary.get("trained_model_warning"),
                    required_tests=[test.name for test in tests] + admet.recommended_followup_tests,
                    batch_priority_score=score,
                    priority_label=label,
                    ranking_reason="Ranked by developability risk, ADMET/Tox concern, Lipinski/Veber status, model status, and data validity.",
                    descriptors=descriptors.model_dump(),
                    drug_likeness=rules.model_dump(),
                    admet_toxicity=admet.model_dump(),
                    model_predictions=model_predictions.model_dump() if model_predictions else None,
                    limitations=[LIMITATION],
                )
            )
        except Exception as exc:
            failed_count += 1
            warnings.append(f"Row {item.row_number} failed during screening: {exc}")
    results.sort(key=lambda row: (-row.batch_priority_score, row.overall_admet_tox_concern_score, row.row_number))
    for index, row in enumerate(results, start=1):
        row.batch_rank = index
    status = model_status_response()
    external_info = next(
        (model for model in status["available_models"] + status["unavailable_models"] if model.model_id == "external_admet_provider_v1"),
        None,
    )
    local_info = next(
        (model for model in status["available_models"] + status["unavailable_models"] if model.model_id == "local_admet_model"),
        None,
    )
    trained_info = next(
        (model for model in status["available_models"] + status["unavailable_models"] if model.model_id == "trained_local_admet_model"),
        None,
    )
    response = BatchLibraryScreenResponse(
        batch_id=batch_id,
        screened_count=len(results),
        failed_count=failed_count,
        results=results,
        ranking_summary={
            "top_compound": results[0].compound_name or results[0].compound_id if results else None,
            "high_priority_count": sum(1 for row in results if row.priority_label == "High"),
            "review_or_low_count": sum(1 for row in results if row.priority_label in {"Review", "Low"}),
        },
        warnings=warnings,
        model_status_summary={
            "available_models": [model.model_id for model in status["available_models"]],
            "unavailable_models": [model.model_id for model in status["unavailable_models"]],
            "external_provider_status": external_info.status if external_info else "not_registered",
            "external_model_available": bool(external_info and external_info.status == "available"),
            "external_model_warning": external_info.warning if external_info else "External ADMET provider adapter is not registered.",
            "local_model_status": local_info.status if local_info else "not_registered",
            "local_model_available": bool(local_info and local_info.status == "available"),
            "local_model_warning": local_info.warning if local_info else "Local ADMET model adapter is not registered.",
            "trained_model_status": trained_info.status if trained_info else "not_registered",
            "trained_model_available": bool(trained_info and trained_info.status == "available"),
            "trained_model_warning": trained_info.warning if trained_info else "Trained local ADMET model adapter is not registered.",
            "prediction_source_used": "Rule-based ADMET/Tox adapter plus unavailable-model messages when requested.",
        },
        limitations=[LIMITATION, "Evidence quality not evaluated because uploaded compounds are not target-linked candidates."],
    )
    response.batch_screening_id = save_batch_library_run(response)
    return response


def save_batch_library_run(response: BatchLibraryScreenResponse) -> int:
    init_db()
    with get_connection() as connection:
        cursor = connection.execute(
            """
            INSERT INTO batch_library_runs (batch_id, screened_count, failed_count, result_payload_json)
            VALUES (?, ?, ?, ?)
            """,
            (response.batch_id, response.screened_count, response.failed_count, response.model_dump_json()),
        )
        return int(cursor.lastrowid)


def get_batch_library_run(run_id: int) -> BatchLibraryRunStored | None:
    init_db()
    with get_connection() as connection:
        row = connection.execute("SELECT * FROM batch_library_runs WHERE id = ?", (run_id,)).fetchone()
    return _stored_run(row) if row else None


def list_batch_library_runs() -> list[BatchLibraryRunStored]:
    init_db()
    with get_connection() as connection:
        rows = connection.execute("SELECT * FROM batch_library_runs ORDER BY datetime(created_at) DESC, id DESC LIMIT 50").fetchall()
    return [_stored_run(row) for row in rows]


def _stored_run(row) -> BatchLibraryRunStored:
    return BatchLibraryRunStored(
        id=row["id"],
        batch_id=row["batch_id"],
        screened_count=row["screened_count"],
        failed_count=row["failed_count"],
        payload=BatchLibraryScreenResponse.model_validate_json(row["result_payload_json"]),
        created_at=row["created_at"],
    )


def batch_library_csv(response: BatchLibraryScreenResponse) -> str:
    headers = [
        "batch_rank", "compound_name", "compound_id", "canonical_smiles", "molecular_weight", "logp", "tpsa",
        "lipinski_pass", "veber_pass", "concern_level", "overall_admet_tox_concern_score", "model_status",
        "decision", "priority_label", "batch_priority_score", "ranking_reason",
    ]
    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=headers, extrasaction="ignore")
    writer.writeheader()
    for row in response.results:
        writer.writerow(row.model_dump())
    return output.getvalue()


def _pdf_table(rows: list[list[str]], widths: list[float] | None = None) -> Table:
    table = Table(rows, colWidths=widths or [1.5 * inch, 5.0 * inch], repeatRows=1 if len(rows) > 1 else 0)
    table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2f5")), ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#c8d1dc")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("FONTSIZE", (0, 0), (-1, -1), 7.2)]))
    return table


def batch_library_pdf(response: BatchLibraryScreenResponse) -> bytes:
    buffer = BytesIO()
    styles = getSampleStyleSheet()
    story = []
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=0.45 * inch, leftMargin=0.45 * inch)
    story.append(Paragraph("DrugScreen360 Batch Compound Library Screening Report", styles["Title"]))
    story.append(Paragraph(DISCLAIMER, styles["BodyText"]))
    story.append(Paragraph(LIMITATION, styles["BodyText"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Uploaded File Summary", styles["Heading2"]))
    story.append(_pdf_table([["Batch ID", str(response.batch_id)], ["Screened", str(response.screened_count)], ["Failed", str(response.failed_count)]]))
    story.append(Paragraph("Screening Method", styles["Heading2"]))
    story.append(Paragraph("RDKit descriptors, Lipinski/Veber, rule-based ADMET/Tox, model registry adapter status, and transparent ranking.", styles["BodyText"]))
    story.append(Paragraph("Ranked Compound Table", styles["Heading2"]))
    rows = [["Rank", "Compound", "MW", "LogP", "TPSA", "ADMET/Tox", "Decision", "Priority"]]
    for row in response.results:
        rows.append([str(row.batch_rank), row.compound_name or row.compound_id or "Unnamed", str(row.molecular_weight), str(row.logp), str(row.tpsa), f"{row.concern_level} ({row.overall_admet_tox_concern_score})", row.decision, row.priority_label])
    story.append(_pdf_table(rows, widths=[0.45 * inch, 1.25 * inch, 0.65 * inch, 0.55 * inch, 0.55 * inch, 0.9 * inch, 1.0 * inch, 0.75 * inch]))
    story.append(Paragraph("Top Candidate Details", styles["Heading2"]))
    for row in response.results[:5]:
        story.append(Paragraph(f"{row.batch_rank}. {row.compound_name or row.compound_id or 'Unnamed'}: {row.ranking_reason}", styles["BodyText"]))
    story.append(Paragraph("Invalid Compounds / Failed Rows", styles["Heading2"]))
    story.append(Paragraph("; ".join(response.warnings) or "No screening failures.", styles["BodyText"]))
    story.append(Paragraph("Prediction Model Status", styles["Heading2"]))
    story.append(_pdf_table([[key.replace("_", " ").title(), str(value)] for key, value in response.model_status_summary.items()]))
    story.append(Paragraph("Required Experimental Follow-up Plan", styles["Heading2"]))
    story.append(Paragraph("Select assays based on compound-specific required tests, ADMET/Tox flags, and expert review.", styles["BodyText"]))
    story.append(Paragraph("Limitations", styles["Heading2"]))
    for item in response.limitations:
        story.append(Paragraph(f"- {item}", styles["BodyText"]))
    doc.build(story)
    return buffer.getvalue()


def batch_library_docx(response: BatchLibraryScreenResponse) -> bytes:
    document = Document()
    document.add_heading("DrugScreen360 Batch Compound Library Screening Report", 0)
    document.add_paragraph(DISCLAIMER)
    document.add_paragraph(LIMITATION)
    document.add_heading("Uploaded File Summary", level=1)
    document.add_paragraph(f"Batch ID: {response.batch_id}; Screened: {response.screened_count}; Failed: {response.failed_count}")
    document.add_heading("Screening Method", level=1)
    document.add_paragraph("RDKit descriptors, Lipinski/Veber, rule-based ADMET/Tox, model registry adapter status, and transparent ranking.")
    document.add_heading("Ranked Compound Table", level=1)
    for row in response.results:
        document.add_paragraph(f"{row.batch_rank}. {row.compound_name or row.compound_id or 'Unnamed'} - {row.priority_label} ({row.batch_priority_score})", style="List Bullet")
    document.add_heading("Top Candidate Details", level=1)
    for row in response.results[:5]:
        document.add_paragraph(f"{row.compound_name or row.compound_id or 'Unnamed'}: {row.ranking_reason}")
    document.add_heading("Invalid Compounds / Failed Rows", level=1)
    for warning in response.warnings or ["No screening failures."]:
        document.add_paragraph(warning, style="List Bullet")
    document.add_heading("Prediction Model Status", level=1)
    for key, value in response.model_status_summary.items():
        document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
    document.add_heading("Required Experimental Follow-up Plan", level=1)
    document.add_paragraph("Select assays based on compound-specific required tests, ADMET/Tox flags, and expert review.")
    document.add_heading("Limitations", level=1)
    for item in response.limitations:
        document.add_paragraph(item, style="List Bullet")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def example_csv() -> str:
    rows = [
        ["smiles", "name", "compound_id", "source", "notes"],
        ["CC(=O)OC1=CC=CC=C1C(=O)O", "Aspirin", "EX001", "example", "reference"],
        ["Cn1cnc2c1c(=O)n(C)c(=O)n2C", "Caffeine", "EX002", "example", "reference"],
        ["CC(C)Cc1ccc(cc1)[C@@H](C)C(=O)O", "Ibuprofen", "EX003", "example", "reference"],
        ["CC(=O)NC1=CC=C(O)C=C1", "Acetaminophen", "EX004", "example", "reference"],
        ["CN(C)C(=N)NC(=N)N", "Metformin", "EX005", "example", "reference"],
    ]
    output = StringIO()
    writer = csv.writer(output)
    writer.writerows(rows)
    return output.getvalue()


def example_smi() -> str:
    return "\n".join([
        "CC(=O)OC1=CC=CC=C1C(=O)O Aspirin",
        "Cn1cnc2c1c(=O)n(C)c(=O)n2C Caffeine",
        "CC(C)Cc1ccc(cc1)[C@@H](C)C(=O)O Ibuprofen",
        "CC(=O)NC1=CC=C(O)C=C1 Acetaminophen",
        "CN(C)C(=N)NC(=N)N Metformin",
    ]) + "\n"
