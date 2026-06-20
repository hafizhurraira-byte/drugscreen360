import base64
from io import BytesIO
from typing import Iterable

from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.models.schemas import ScreeningReport


def _image_bytes(data_url: str | None) -> BytesIO | None:
    if not data_url or "," not in data_url:
        return None
    _, encoded = data_url.split(",", 1)
    return BytesIO(base64.b64decode(encoded))


def _pairs(mapping: dict) -> list[list[str]]:
    return [[str(key).replace("_", " ").title(), str(value)] for key, value in mapping.items()]


def _pdf_table(rows: list[list[str]]) -> Table:
    table = Table(rows, colWidths=[2.1 * inch, 4.4 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef2f5")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#c8d1dc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def _bullet_text(items: Iterable[str]) -> str:
    return "<br/>".join(f"- {item}" for item in items)


def _admet_tox_rows(report: ScreeningReport) -> list[list[str]]:
    assessment = report.admet_toxicity_v1
    if assessment is None:
        return [["Status", "Not available"]]
    return [
        ["Absorption Risk", assessment.absorption.absorption_risk],
        ["Solubility Risk", assessment.solubility.solubility_risk],
        ["BBB/CNS Flag", assessment.bbb_cns_flag.bbb_exposure_flag],
        ["Metabolism/CYP Status", assessment.metabolism_status.cyp_prediction_status],
        ["Structural Alert Risk", assessment.structural_alerts.structural_alert_risk],
        ["Structural Alerts", ", ".join(assessment.structural_alerts.structural_alerts) or "None found"],
        ["hERG Status", assessment.herg_status.prediction_status],
        ["Genotoxicity Status", assessment.ames_genotoxicity_status.prediction_status],
        ["Hepatotoxicity Status", assessment.hepatotoxicity_status.prediction_status],
        ["Overall Concern Score", str(assessment.overall.overall_admet_tox_concern_score)],
        ["Concern Level", assessment.overall.concern_level],
        ["Confidence Level", assessment.overall.confidence_level],
        ["Recommended Follow-ups", "; ".join(assessment.recommended_followup_tests)],
        ["Limitations", "; ".join(assessment.limitations)],
    ]


def _evidence_rows(report: ScreeningReport) -> list[list[str]]:
    evidence = report.evidence_quality
    if evidence is None:
        return [
            [
                "Status",
                "Evidence quality not evaluated because this screening was run from compound identity only, not from a target-linked candidate record.",
            ]
        ]
    return [
        ["Evidence Score", str(evidence.evidence_score)],
        ["Evidence Level", evidence.evidence_level],
        ["Potency Quality", evidence.potency_quality],
        ["Data Quality Score", str(evidence.data_quality_score)],
        ["Target Confidence", evidence.target_confidence_summary],
        ["Reasons", "; ".join(evidence.evidence_reasons)],
        ["Warnings", "; ".join(evidence.warnings) or "None"],
        ["Recommended Next Step", evidence.recommended_action],
        ["BindingDB Status", evidence.bindingdb_support.limitation],
        ["Limitation", evidence.limitation],
    ]


def _model_prediction_rows(report: ScreeningReport) -> list[list[str]]:
    predictions = report.model_predictions
    if predictions is None:
        return [["Status", "No model prediction bundle was saved with this report."]]
    rows = [["Model", "Status / Source / Confidence"]]
    for bundle in predictions.model_outputs:
        rows.append([bundle.model_name, f"{bundle.model_status} / {bundle.prediction_source} / {bundle.confidence}"])
    rows.append(["Interpretation", predictions.combined_interpretation])
    rows.append(["Warnings", "; ".join(predictions.warnings) or "None"])
    return rows


def build_pdf_report(report: ScreeningReport) -> bytes:
    buffer = BytesIO()
    styles = getSampleStyleSheet()
    story = []

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )

    story.append(Paragraph("DrugScreen360 Candidate Screening Report", styles["Title"]))
    story.append(Paragraph(report.disclaimer, styles["BodyText"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Compound Identity", styles["Heading2"]))
    identity = report.compound_identity
    story.append(
        _pdf_table(
            [
                ["Name", identity.compound_name or "Not available"],
                ["PubChem CID", str(identity.pubchem_cid or "Not available")],
                ["Formula", identity.molecular_formula or "Not available"],
                ["Molecular Weight", str(identity.molecular_weight or "Not available")],
                ["IUPAC Name", identity.iupac_name or "Not available"],
                ["Canonical SMILES", identity.canonical_smiles or "Not available"],
                ["PubChem Link", identity.pubchem_source_link or "Not available"],
            ]
        )
    )
    image_data = _image_bytes(identity.structure_image_base64)
    if image_data:
        story.append(Spacer(1, 10))
        story.append(Image(image_data, width=3.6 * inch, height=2.5 * inch))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Physicochemical Properties", styles["Heading2"]))
    story.append(_pdf_table(_pairs(report.physicochemical_properties.model_dump())))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Drug-Likeness Assessment", styles["Heading2"]))
    story.append(_pdf_table(_pairs(report.drug_likeness.model_dump())))
    story.append(Spacer(1, 12))

    story.append(Paragraph("ADMET Placeholder Section", styles["Heading2"]))
    story.append(Paragraph(report.admet_placeholder.message, styles["BodyText"]))
    story.append(Paragraph(_bullet_text(report.admet_placeholder.future_outputs), styles["BodyText"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Toxicity Placeholder Section", styles["Heading2"]))
    story.append(Paragraph(report.toxicity_placeholder.message, styles["BodyText"]))
    story.append(Paragraph(_bullet_text(report.toxicity_placeholder.future_outputs), styles["BodyText"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("ADMET/Toxicity Rule-Based Assessment", styles["Heading2"]))
    if report.admet_toxicity_v1:
        story.append(Paragraph(report.admet_toxicity_v1.label, styles["BodyText"]))
    story.append(_pdf_table(_admet_tox_rows(report)))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Evidence Quality Assessment", styles["Heading2"]))
    story.append(_pdf_table(_evidence_rows(report)))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Prediction Model Status", styles["Heading2"]))
    story.append(_pdf_table(_model_prediction_rows(report)))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Required Experimental Tests", styles["Heading2"]))
    story.append(
        _pdf_table(
            [["Test", "Priority / Reason"]]
            + [[test.name, f"{test.priority}: {test.reason}"] for test in report.required_lab_tests]
        )
    )
    story.append(Spacer(1, 12))

    story.append(Paragraph("Go / No-Go Recommendation", styles["Heading2"]))
    story.append(_pdf_table(_pairs(report.go_no_go_recommendation)))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Limitations", styles["Heading2"]))
    story.append(Paragraph(_bullet_text(report.limitations), styles["BodyText"]))

    doc.build(story)
    return buffer.getvalue()


def build_docx_report(report: ScreeningReport) -> bytes:
    document = Document()
    document.add_heading("DrugScreen360 Candidate Screening Report", 0)
    document.add_paragraph(report.disclaimer)

    document.add_heading("Compound Identity", level=1)
    identity = report.compound_identity
    for label, value in [
        ("Name", identity.compound_name),
        ("PubChem CID", identity.pubchem_cid),
        ("Formula", identity.molecular_formula),
        ("Molecular Weight", identity.molecular_weight),
        ("IUPAC Name", identity.iupac_name),
        ("Canonical SMILES", identity.canonical_smiles),
        ("PubChem Link", identity.pubchem_source_link),
    ]:
        document.add_paragraph(f"{label}: {value or 'Not available'}")

    image_data = _image_bytes(identity.structure_image_base64)
    if image_data:
        document.add_picture(image_data, width=Inches(4.8))

    document.add_heading("Physicochemical Properties", level=1)
    for key, value in report.physicochemical_properties.model_dump().items():
        document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

    document.add_heading("Drug-Likeness Assessment", level=1)
    for key, value in report.drug_likeness.model_dump().items():
        document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

    document.add_heading("ADMET Placeholder Section", level=1)
    document.add_paragraph(report.admet_placeholder.message)
    for item in report.admet_placeholder.future_outputs:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Toxicity Placeholder Section", level=1)
    document.add_paragraph(report.toxicity_placeholder.message)
    for item in report.toxicity_placeholder.future_outputs:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("ADMET/Toxicity Rule-Based Assessment", level=1)
    if report.admet_toxicity_v1:
        document.add_paragraph(report.admet_toxicity_v1.label)
    for label, value in _admet_tox_rows(report):
        document.add_paragraph(f"{label}: {value}")

    document.add_heading("Evidence Quality Assessment", level=1)
    for label, value in _evidence_rows(report):
        document.add_paragraph(f"{label}: {value}")

    document.add_heading("Prediction Model Status", level=1)
    for label, value in _model_prediction_rows(report):
        document.add_paragraph(f"{label}: {value}")

    document.add_heading("Required Experimental Tests", level=1)
    for test in report.required_lab_tests:
        document.add_paragraph(f"{test.name} ({test.priority}): {test.reason}", style="List Bullet")

    document.add_heading("Go / No-Go Recommendation", level=1)
    for key, value in report.go_no_go_recommendation.items():
        document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

    document.add_heading("Limitations", level=1)
    for item in report.limitations:
        document.add_paragraph(item, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
