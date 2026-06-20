import json
import platform
import sys
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import HTTPException
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.database import get_connection, init_db
from app.models.project_workspace_report_models import (
    ProjectWorkspaceReportCreateRequest,
    ProjectWorkspaceReportCreateResponse,
    ProjectWorkspaceReportListItem,
)
from app.services.local_admet_model import validate_local_admet_model
from app.models.project_workspace_models import ProjectAttachRequest
from app.services.project_workspace_service import attach_project_item, project_dashboard
from app.services.version import app_version

REPORT_DIR = Path(__file__).resolve().parents[2] / "project_workspace_reports"

DISCLAIMER = (
    "Computational decision-support only. This report does not prove safety, efficacy, "
    "clinical success, regulatory approval, or market readiness. Laboratory validation and expert review are required."
)


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _timestamp(dt: datetime) -> str:
    return dt.strftime("%Y-%m-%d_%H-%M-%S")


def _safe(value: Any) -> str:
    if value in (None, "", [], {}):
        return "not available"
    if isinstance(value, list):
        return "; ".join(str(item) for item in value) or "not available"
    if isinstance(value, dict):
        return json.dumps(value, default=str)
    return str(value)


def _pairs(mapping: dict[str, Any]) -> list[list[str]]:
    return [[key.replace("_", " ").title(), _safe(value)] for key, value in mapping.items()]


def _pdf_table(rows: list[list[str]], widths: list[float] | None = None) -> Table:
    table = Table(rows, colWidths=widths or [2.1 * inch, 4.4 * inch], repeatRows=1 if len(rows) > 1 else 0)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2f5")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#c8d1dc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.2),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _matrix_rows(dashboard: dict[str, Any]) -> list[list[str]]:
    rows = [["Candidate", "Workflow", "Target", "MW", "LogP", "TPSA", "ADMET/Tox", "Evidence", "Model", "Decision"]]
    for row in dashboard.get("candidate_matrix") or []:
        rows.append(
            [
                _safe(row.get("candidate_name")),
                _safe(row.get("source_workflow")),
                _safe(row.get("target_name")),
                _safe(row.get("molecular_weight")),
                _safe(row.get("logp")),
                _safe(row.get("tpsa")),
                _safe(row.get("admet_risk_summary")),
                f"{_safe(row.get('evidence_level'))} ({_safe(row.get('evidence_score'))})",
                _safe(row.get("model_prediction_status")),
                _safe(row.get("decision_label")),
            ]
        )
    return rows


def _report_payload(project_id: int, options: ProjectWorkspaceReportCreateRequest, created_at: str) -> dict[str, Any]:
    dashboard = project_dashboard(project_id).model_dump()
    local_validation = validate_local_admet_model()
    project = dashboard["project"]
    warnings = list(dashboard.get("warnings") or [])
    if not dashboard.get("candidate_matrix"):
        warnings.append("Candidate decision matrix is empty because no candidate-level data is attached.")
    payload = {
        "title": f"DrugScreen360 Project Workspace Report - {project['title']}",
        "created_at": created_at,
        "app_version": app_version(),
        "disclaimer": DISCLAIMER,
        "project": project,
        "dashboard_summary": dashboard.get("summary_cards", {}),
        "item_counts": dashboard.get("item_counts", {}),
        "risk_summary": dashboard.get("risk_summary", {}),
        "recommended_next_steps": dashboard.get("recommended_next_steps", []),
        "candidate_matrix": dashboard.get("candidate_matrix", []) if options.include_candidate_matrix else [],
        "model_status_summary": dashboard.get("model_status_summary", {}) if options.include_model_status else {"status": "not included"},
        "local_model_validation": local_validation if options.include_model_status else {"status": "not included"},
        "limitations": dashboard.get("limitations", []) if options.include_limitations else [],
        "reproducibility": {
            "app_version": app_version(),
            "python_version": sys.version.split()[0],
            "python_implementation": platform.python_implementation(),
            "created_at": created_at,
            "run_command": ".\\scripts\\start_all.ps1",
            "test_command": ".\\scripts\\run_tests.ps1",
        }
        if options.include_reproducibility
        else {"status": "not included"},
        "warnings": warnings,
    }
    return payload


def _build_pdf(payload: dict[str, Any]) -> bytes:
    buffer = BytesIO()
    styles = getSampleStyleSheet()
    story = []
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=0.45 * inch, leftMargin=0.45 * inch)
    project = payload["project"]
    story.append(Paragraph(payload["title"], styles["Title"]))
    story.append(Paragraph(payload["disclaimer"], styles["BodyText"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Project Metadata", styles["Heading2"]))
    story.append(
        _pdf_table(
            _pairs(
                {
                    "project_id": project.get("id"),
                    "title": project.get("title"),
                    "status": project.get("status"),
                    "project_type": project.get("project_type"),
                    "disease_area": project.get("disease_area"),
                    "target_name": project.get("target_name"),
                    "description": project.get("description"),
                    "notes": project.get("notes"),
                }
            )
        )
    )
    story.append(Spacer(1, 10))
    story.append(Paragraph("Project Dashboard Summary", styles["Heading2"]))
    story.append(_pdf_table(_pairs(payload["dashboard_summary"])))
    story.append(Paragraph("Attached Item Summary", styles["Heading2"]))
    story.append(_pdf_table(_pairs(payload["item_counts"] or {"items": "not available"})))
    story.append(Paragraph("Risk And Evidence Summary", styles["Heading2"]))
    story.append(_pdf_table(_pairs(payload["risk_summary"])))
    if payload["candidate_matrix"]:
        story.append(Paragraph("Candidate Decision Matrix", styles["Heading2"]))
        story.append(_pdf_table(_matrix_rows(payload), widths=[0.8 * inch, 0.75 * inch, 0.65 * inch, 0.45 * inch, 0.45 * inch, 0.45 * inch, 0.7 * inch, 0.75 * inch, 0.75 * inch, 1.1 * inch]))
    else:
        story.append(Paragraph("Candidate Decision Matrix", styles["Heading2"]))
        story.append(Paragraph("No candidate-level data available for this project.", styles["BodyText"]))
    story.append(Paragraph("Model Status Summary", styles["Heading2"]))
    story.append(_pdf_table(_pairs(payload["model_status_summary"])))
    story.append(Paragraph("Local Model Validation Summary", styles["Heading2"]))
    story.append(_pdf_table(_pairs(payload["local_model_validation"])))
    story.append(Paragraph("Recommended Next Steps", styles["Heading2"]))
    for item in payload["recommended_next_steps"]:
        story.append(Paragraph(f"- {item}", styles["BodyText"]))
    story.append(Paragraph("Scientific Limitations", styles["Heading2"]))
    for item in payload["limitations"]:
        story.append(Paragraph(f"- {item}", styles["BodyText"]))
    story.append(Paragraph("Reproducibility", styles["Heading2"]))
    story.append(_pdf_table(_pairs(payload["reproducibility"])))
    story.append(Paragraph("Warnings", styles["Heading2"]))
    for item in payload["warnings"] or ["None"]:
        story.append(Paragraph(f"- {item}", styles["BodyText"]))
    doc.build(story)
    return buffer.getvalue()


def _build_docx(payload: dict[str, Any]) -> bytes:
    document = Document()
    project = payload["project"]
    document.add_heading(payload["title"], 0)
    document.add_paragraph(payload["disclaimer"])
    document.add_heading("Project Metadata", level=1)
    for key, value in {
        "Project ID": project.get("id"),
        "Title": project.get("title"),
        "Status": project.get("status"),
        "Project type": project.get("project_type"),
        "Disease area": project.get("disease_area"),
        "Target name": project.get("target_name"),
        "Description": project.get("description"),
        "Notes": project.get("notes"),
    }.items():
        document.add_paragraph(f"{key}: {_safe(value)}")
    for heading, mapping in [
        ("Project Dashboard Summary", payload["dashboard_summary"]),
        ("Attached Item Summary", payload["item_counts"] or {"items": "not available"}),
        ("Risk And Evidence Summary", payload["risk_summary"]),
        ("Model Status Summary", payload["model_status_summary"]),
        ("Local Model Validation Summary", payload["local_model_validation"]),
        ("Reproducibility", payload["reproducibility"]),
    ]:
        document.add_heading(heading, level=1)
        for key, value in mapping.items():
            document.add_paragraph(f"{key.replace('_', ' ').title()}: {_safe(value)}")
    document.add_heading("Candidate Decision Matrix", level=1)
    if payload["candidate_matrix"]:
        for row in payload["candidate_matrix"]:
            document.add_paragraph(
                f"{_safe(row.get('candidate_name'))}: {_safe(row.get('decision_label'))}; "
                f"ADMET/Tox {_safe(row.get('admet_risk_summary'))}; Evidence {_safe(row.get('evidence_level'))}; "
                f"Model {_safe(row.get('model_prediction_status'))}.",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No candidate-level data available for this project.")
    document.add_heading("Recommended Next Steps", level=1)
    for item in payload["recommended_next_steps"]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("Scientific Limitations", level=1)
    for item in payload["limitations"]:
        document.add_paragraph(item, style="List Bullet")
    document.add_heading("Warnings", level=1)
    for item in payload["warnings"] or ["None"]:
        document.add_paragraph(item, style="List Bullet")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def create_project_workspace_report(project_id: int, options: ProjectWorkspaceReportCreateRequest) -> ProjectWorkspaceReportCreateResponse:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    created = _now()
    created_at = created.isoformat()
    stamp = _timestamp(created)
    payload = _report_payload(project_id, options, created_at)
    base = f"project_{project_id}_workspace_report_{stamp}"
    pdf_name = f"{base}.pdf"
    docx_name = f"{base}.docx"
    json_name = f"{base}.json"
    (REPORT_DIR / pdf_name).write_bytes(_build_pdf(payload))
    (REPORT_DIR / docx_name).write_bytes(_build_docx(payload))
    (REPORT_DIR / json_name).write_text(json.dumps(payload, indent=2, default=str), encoding="utf-8")
    init_db()
    with get_connection() as connection:
        cursor = connection.execute(
            """
            INSERT INTO project_workspace_reports (project_id, filename_pdf, filename_docx, filename_json, warnings_json)
            VALUES (?, ?, ?, ?, ?)
            """,
            (project_id, pdf_name, docx_name, json_name, json.dumps(payload["warnings"])),
        )
        report_id = int(cursor.lastrowid)
        row = connection.execute("SELECT created_at FROM project_workspace_reports WHERE id = ?", (report_id,)).fetchone()
    attach_project_item(
        project_id,
        ProjectAttachRequest(
            item_type="project_workspace_report",
            item_id=str(report_id),
            item_title=payload["title"],
            metadata={
                "workflow_type": "project_workspace_report",
                "project_id": project_id,
                "report_id": report_id,
                "decision": "not evaluated",
                "model_status": payload.get("model_status_summary"),
                "created_at": row["created_at"] if row else created_at,
            },
        ),
    )
    return ProjectWorkspaceReportCreateResponse(
        report_id=report_id,
        project_id=project_id,
        created_at=row["created_at"] if row else created_at,
        pdf_url=f"/api/projects/{project_id}/report/{report_id}/pdf",
        docx_url=f"/api/projects/{project_id}/report/{report_id}/docx",
        json_url=f"/api/projects/{project_id}/report/{report_id}/json",
        warnings=payload["warnings"],
    )


def list_project_workspace_reports(project_id: int) -> list[ProjectWorkspaceReportListItem]:
    project_dashboard(project_id)
    init_db()
    with get_connection() as connection:
        rows = connection.execute(
            "SELECT * FROM project_workspace_reports WHERE project_id = ? ORDER BY datetime(created_at) DESC, id DESC",
            (project_id,),
        ).fetchall()
    return [
        ProjectWorkspaceReportListItem(
            report_id=row["id"],
            project_id=row["project_id"],
            created_at=row["created_at"],
            pdf_url=f"/api/projects/{project_id}/report/{row['id']}/pdf",
            docx_url=f"/api/projects/{project_id}/report/{row['id']}/docx",
            json_url=f"/api/projects/{project_id}/report/{row['id']}/json",
            warnings=json.loads(row["warnings_json"]),
            filename_pdf=row["filename_pdf"],
            filename_docx=row["filename_docx"],
            filename_json=row["filename_json"],
        )
        for row in rows
    ]


def get_project_workspace_report_path(project_id: int, report_id: int, file_format: str) -> Path:
    if file_format not in {"pdf", "docx", "json"}:
        raise HTTPException(status_code=404, detail="Unsupported report format.")
    project_dashboard(project_id)
    column = f"filename_{file_format}"
    init_db()
    with get_connection() as connection:
        row = connection.execute(
            f"SELECT {column} FROM project_workspace_reports WHERE id = ? AND project_id = ?",
            (report_id, project_id),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Project workspace report not found.")
    path = REPORT_DIR / row[column]
    if not path.exists():
        raise HTTPException(status_code=404, detail="Project workspace report file is missing.")
    return path
