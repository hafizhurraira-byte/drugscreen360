import json
import uuid
import zipfile
from io import BytesIO

from fastapi.testclient import TestClient

from app.database import get_connection, init_db
from app.main import app
from app.services import final_report_service, research_export_service
from app.services.disease_to_lead_context import save_disease_to_lead_run

client = TestClient(app)


def _create_final(**extra):
    payload = {
        "report_title": "DrugScreen360 Final Project Report",
        "include_screening": True,
        "include_admet_prediction": True,
        "include_model_training": True,
        "include_external_validation": True,
        "include_applicability_domain": True,
        "include_explainability": True,
        "include_lead_prioritization": True,
        "include_validation_planner": True,
        "include_experimental_feedback": True,
        "formats": ["json", "pdf", "docx"],
    }
    payload.update(extra)
    return client.post("/api/final-report/create", json=payload)


def _save_disease_to_lead_snapshot(project_id, disease, target, known_compound, candidates):
    return save_disease_to_lead_run(
        {
            "workflow_id": f"test-{uuid.uuid4()}",
            "project_id": project_id,
            "report_id": None,
            "disease_name_raw": disease,
            "disease_name_normalized": disease,
            "user_entered_target_raw": target,
            "user_entered_target_normalized": target,
            "resolved_target_name": target,
            "resolved_target_id": target,
            "resolved_target_gene_symbol": target,
            "resolved_target_organism": "Homo sapiens",
            "target_resolution_confidence": 100.0,
            "target_resolution_status": "exact_symbol_match",
            "known_compound_raw": known_compound,
            "known_compound_normalized": known_compound,
            "known_compound_id": f"KNOWN-{known_compound.upper()}",
            "candidate_limit": 5,
            "similarity_limit": 5,
            "analysis_depth": "quick",
            "scoring_profile": "balanced_admet",
            "generated_candidate_list": candidates,
            "deduplicated_candidate_list": candidates,
            "duplicate_records_removed": 0,
            "admet_results": [],
            "prioritization_results": {"ranked_candidates": candidates, "warnings": []},
            "validation_planner_results": {},
            "missing_evidence_summary": [],
        }
    )


def test_final_report_creation_works_with_minimal_data(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    response = _create_final()
    assert response.status_code == 200
    body = response.json()
    assert body["report_id"]
    assert body["generated_files"]["json"].endswith("/json")
    assert "Computational decision-support report only" in body["scientific_notice"]
    assert body["missing_sections"] is not None


def test_final_report_json_includes_evidence_package_provenance(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    project = client.post("/api/projects/create", json={"title": "M2 Evidence Project", "project_type": "disease_screening", "status": "active"}).json()
    _save_disease_to_lead_snapshot(
        project["id"],
        "non-small cell lung cancer",
        "EGFR",
        "Erlotinib",
        [{"compound_name": "Erlotinib", "compound_id": "KNOWN-ERLOTINIB", "smiles": "C#Cc1cccc(Nc2ncnc3cc(OCCOC)c(OCCOC)cc23)c1", "total_score": 0.72}],
    )

    report_id = _create_final(
        project_id=project["id"],
        report_title="M2 Evidence Package Report",
        report_mode="concise_disease_to_lead_report",
    ).json()["report_id"]
    report = client.get(f"/api/final-report/reports/{report_id}/json").json()

    assert "evidence_package" in report
    assert "MODEL PREDICTION" in report["evidence_package"]["evidence_type_definitions"]
    assert report["model_evidence_list"][0]["evidence_type"] == "RULE-BASED HEURISTIC"
    assert report["model_evidence_list"][0]["source"] == "DrugScreen360 descriptor and rule-based workflow"


def test_final_report_includes_egfr_activity_model_prediction(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    project = client.post("/api/projects/create", json={"title": "EGFR Activity Report", "project_type": "disease_screening", "status": "active"}).json()
    run_id = _save_disease_to_lead_snapshot(
        project["id"],
        "non-small cell lung cancer",
        "EGFR",
        "Erlotinib",
        [
            {
                "compound_name": "Erlotinib",
                "compound_id": "CHEMBL553",
                "smiles": "C#Cc1cccc(Nc2ncnc3cc(OCCOC)c(OCCOC)cc23)c1",
                "total_score": 91.0,
                "priority_label": "high_priority_for_review",
                "activity_model_prediction": {
                    "status": "available",
                    "model_id": "egfr_activity_v2",
                    "model_name": "random_forest_180_morgan",
                    "model_version": "v2",
                    "predicted_pIC50": 7.1,
                    "predicted_IC50_nM": 79.4,
                    "dataset_lineage": {"sources": ["ChEMBL EGFR curated v2", "BindingDB augmentation subset"]},
                    "validation_status": "externally_validated_research_use",
                    "applicability_domain_status": "IN_DOMAIN",
                    "uncertainty_value": 0.5,
                    "artifact_hash": "7bd850e41d877a0d3c1c39dde42914ba67fa81142962c7ca7e67d7707f1b6c61",
                    "nearest_training_similarity": 0.8,
                    "interval_lower": 5.9,
                    "interval_upper": 8.3,
                    "external_observed_coverage": 0.8337,
                },
            }
        ],
    )
    response = _create_final(
        project_id=project["id"],
        report_mode="concise_disease_to_lead_report",
        disease_to_lead_run_id=run_id,
        formats=["json"],
    )
    report = client.get(response.json()["generated_files"]["json"]).json()
    evidence = next(item for item in report["model_evidence_list"] if item["active_model_id"] == "egfr_activity_v2")

    assert evidence["source"] == "DrugScreen360 EGFR target-specific activity model"
    assert evidence["endpoint_predicted"] == "EGFR IC50/pIC50"
    assert "model-derived IC50_nM" in evidence["prediction"]
    assert evidence["external_observed_coverage"] == 0.8337


def test_final_report_creation_works_with_project_data(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    project = client.post("/api/projects/create", json={"title": "Final Report Project", "project_type": "general_research", "status": "active"}).json()
    lead = client.post(
        "/api/admet-leads/prioritize",
        json={
            "project_id": project["id"],
            "source_type": "manual",
            "candidates": [{"compound_name": "Ethanol", "smiles": "CCO"}],
            "include_trained_model": False,
            "include_domain": False,
            "include_explainability": False,
        },
    )
    assert lead.status_code == 200
    response = _create_final(project_id=project["id"])
    assert response.status_code == 200
    body = response.json()
    assert body["project_id"] == project["id"]
    assert "lead_prioritization" in body["included_sections"]


def test_final_report_missing_sections_are_reported(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    response = _create_final(
        include_external_validation=True,
        include_applicability_domain=True,
        include_explainability=True,
    )
    assert response.status_code == 200
    body = response.json()
    assert isinstance(body["missing_sections"], list)


def test_json_pdf_docx_endpoints_work(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    report_id = _create_final().json()["report_id"]
    json_response = client.get(f"/api/final-report/reports/{report_id}/json")
    pdf_response = client.get(f"/api/final-report/reports/{report_id}/pdf")
    docx_response = client.get(f"/api/final-report/reports/{report_id}/docx")
    assert json_response.status_code == 200
    assert pdf_response.status_code == 200
    assert docx_response.status_code == 200
    assert b"DrugScreen360 Final Project Report" in json_response.content
    assert pdf_response.headers["content-type"].startswith("application/pdf")
    assert "wordprocessingml" in docx_response.headers["content-type"]


def test_no_fake_experimental_results_or_clinical_claims(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    report_id = _create_final().json()["report_id"]
    report = client.get(f"/api/final-report/reports/{report_id}/json").json()
    serialized = str(report).lower()
    assert "confirmed safe" not in serialized
    assert "drug approved" not in serialized
    assert "clinically validated" not in serialized
    assert "simulated assay result" not in serialized
    assert "no experimental results are generated" in serialized


def test_project_attachment_works(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    project = client.post("/api/projects/create", json={"title": "Attached Final Report", "project_type": "general_research", "status": "active"}).json()
    response = _create_final(project_id=project["id"])
    assert response.status_code == 200
    detail = client.get(f"/api/projects/{project['id']}").json()
    assert any(item["item_type"] == "final_project_report" for item in detail["items"])


def test_research_export_includes_final_report_files(tmp_path, monkeypatch):
    final_dir = tmp_path / "final_project_reports"
    export_dir = tmp_path / "research_exports"
    monkeypatch.setattr(final_report_service, "REPORT_DIR", final_dir)
    monkeypatch.setattr(research_export_service, "EXPORT_DIR", export_dir)
    response = _create_final()
    assert response.status_code == 200
    export = client.post(
        "/api/research-export/create",
        json={
            "include_screening_history": False,
            "include_benchmark_runs": False,
            "include_batch_runs": False,
            "include_cache_status": False,
            "include_reports": False,
        },
    )
    assert export.status_code == 200
    download = client.get(export.json()["download_url"])
    assert download.status_code == 200
    with zipfile.ZipFile(BytesIO(download.content)) as archive:
        names = archive.namelist()
    assert any("FINAL_PROJECT_REPORTS/" in name for name in names)
    assert any(name.endswith(".json") and "final_project_report" in name for name in names)


def test_report_list_endpoint_works(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    report_id = _create_final().json()["report_id"]
    response = client.get("/api/final-report/reports")
    assert response.status_code == 200
    assert any(item["report_id"] == report_id for item in response.json())


def test_concise_disease_to_lead_report_quality(tmp_path, monkeypatch):
    from docx import Document

    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    monkeypatch.setattr(final_report_service, "get_active_trained_model_info", lambda: {"status": "unavailable"})
    monkeypatch.setattr("app.services.admet_lead_service.predict_admet_endpoints", lambda smiles, endpoints=None: {"results": []})
    
    # 1. Create a project
    project = client.post("/api/projects/create", json={"title": "EGFR Lead Optimization Project", "project_type": "general_research", "status": "active"}).json()
    project_id = project["id"]
    
    # 2. Run prioritization to populate candidates
    lead = client.post(
        "/api/admet-leads/prioritize",
        json={
            "project_id": project_id,
            "source_type": "manual",
            "candidates": [
                {"compound_name": "Aspirin", "smiles": "CC(=O)Oc1ccccc1C(=O)O"},
                {"compound_name": "Ibuprofen", "smiles": "CC(C)Cc1ccc(cc1)C(C)C(=O)O"}
            ],
            "include_trained_model": False,
            "include_domain": False,
            "include_explainability": False,
        },
    )
    assert lead.status_code == 200
    run_id = lead.json()["run_id"]
    
    # 3. Create a validation plan
    plan = client.post(
        "/api/validation-planner/create",
        json={
            "project_id": project_id,
            "source_type": "manual",
            "plan_title": "Validation Plan: EGFR Test",
            "candidates": [
                {"compound_name": "Aspirin", "smiles": "CC(=O)Oc1ccccc1C(=O)O", "compound_id": "Aspirin"}
            ]
        }
    )
    assert plan.status_code == 200
    plan_id = plan.json()["plan_id"]
    
    # 4. Generate the final report in concise mode
    report_response = client.post(
        "/api/final-report/create",
        json={
            "project_id": project_id,
            "report_title": "Disease-to-Lead Final Concise Report",
            "report_mode": "concise_disease_to_lead_report",
            "prioritization_run_id": run_id,
            "validation_plan_id": plan_id,
            "formats": ["json", "pdf", "docx"],
        }
    )
    assert report_response.status_code == 200
    report_id = report_response.json()["report_id"]
    
    # 5. Download the DOCX via the real download endpoint
    download_response = client.get(f"/api/final-report/reports/{report_id}/docx")
    assert download_response.status_code == 200
    assert "wordprocessingml" in download_response.headers["content-type"]
    
    # 6. Open with python-docx and scan content
    doc_bytes = BytesIO(download_response.content)
    doc = Document(doc_bytes)
    
    # Extract all text from paragraphs and tables
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    combined_text = "\n".join(full_text)
    
    # Assertions for 13 clean sections (checking presence of section headers / key text)
    expected_sections = [
        "Executive Summary",
        "Workflow Input Table",
        "Workflow Completion Table",
        "Top Candidate Table",
        "Top Candidate Interpretation",
        "ADMET & Drug-likeness Summary",
        "Model Evidence & Prediction Confidence",
        "External Validation Summary",
        "Experimental Feedback Summary",
        "Validation Planner Summary",
        "Limitations",
        "Recommended Next Steps",
        "Reproducibility"
    ]
    for section in expected_sections:
        assert section in combined_text, f"Missing expected section: {section}"
        
    # Assertions for forbidden strings
    forbidden_strings = [
        "Top Candidates: [{",
        '{"rank"',
        '"compound_name"',
        '"score_components"',
        '"descriptors"',
        "Stored records summarized",
        "Latest Training Runs: [{",
        "Available Models: [{",
        "Unavailable Models: [{",
        "Missing sections: none"
    ]
    for forbidden in forbidden_strings:
        assert forbidden not in combined_text, f"Found forbidden string: '{forbidden}'"
        
    # Assertions for fallbacks (since experimental results don't exist, and active model is likely unavailable)
    assert "External validation/calibration was not available for this project." in combined_text
    assert "No user-entered experimental assay results were imported. Experimental feedback comparison was not performed." in combined_text
    assert "No active compatible trained ADMET model was available for this run. Ranking used descriptor-based and rule-based evidence only." in combined_text


def test_final_report_polish_requirements(tmp_path, monkeypatch):
    from docx import Document
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    
    # 1. Create a project
    project = client.post("/api/projects/create", json={"title": "EGFR Polish Project", "project_type": "general_research", "status": "active"}).json()
    project_id = project["id"]
    
    # 2. Run prioritization with duplicate compounds
    lead = client.post(
        "/api/admet-leads/prioritize",
        json={
            "project_id": project_id,
            "source_type": "manual",
            "candidates": [
                {"compound_name": "Aspirin", "smiles": "CC(=O)Oc1ccccc1C(=O)O"},
                {"compound_name": "ASPIRIN", "smiles": "CC(=O)Oc1ccccc1C(=O)O"},
                {"compound_name": "Ibuprofen", "smiles": "CC(C)Cc1ccc(cc1)C(C)C(=O)O"}
            ],
            "include_trained_model": False,
            "include_domain": False,
            "include_explainability": False,
        },
    )
    assert lead.status_code == 200
    run_id = lead.json()["run_id"]
    
    # 3. Create a validation plan
    plan = client.post(
        "/api/validation-planner/create",
        json={
            "project_id": project_id,
            "source_type": "manual",
            "plan_title": "Validation Plan: Polish Test",
            "candidates": [
                {"compound_name": "Aspirin", "smiles": "CC(=O)Oc1ccccc1C(=O)O", "compound_id": "Aspirin"}
            ]
        }
    )
    assert plan.status_code == 200
    plan_id = plan.json()["plan_id"]
    
    # 4. Generate the final report
    report_response = client.post(
        "/api/final-report/create",
        json={
            "project_id": project_id,
            "report_title": "Concise Disease-to-Lead Quality Polish Report",
            "report_mode": "concise_disease_to_lead_report",
            "prioritization_run_id": run_id,
            "validation_plan_id": plan_id,
            "disease_name": "Breast Cancer",
            "user_entered_target": "EGFR",
            "resolved_target": "Epidermal growth factor receptor",
            "known_compound": "Aspirin",
            "formats": ["json", "pdf", "docx"],
        }
    )
    assert report_response.status_code == 200
    report_id = report_response.json()["report_id"]
    
    # 5. Download the DOCX
    download_response = client.get(f"/api/final-report/reports/{report_id}/docx")
    assert download_response.status_code == 200
    
    # 6. Open with python-docx and scan content
    doc_bytes = BytesIO(download_response.content)
    doc = Document(doc_bytes)
    
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    combined_text = "\n".join(full_text)
    
    assert "EGFR" in combined_text
    assert "Epidermal growth factor receptor" in combined_text
    # Equivalent target should not trigger warning but the expansions notice
    assert "Resolved target expands or matches the user-entered target." in combined_text
    assert "removing 1 duplicate record(s)" in combined_text
    
    assert "MW" in combined_text
    assert "LogP" in combined_text
    assert "TPSA" in combined_text
    
    # Check that actual numerical values appear in the text/tables
    assert any("180." in text or "206." in text for text in full_text)
    assert "Category:" in combined_text
    
    assert "Rationale None" not in combined_text
    assert "Rationale: None" not in combined_text
    assert "Rationale: null" not in combined_text
    
    # Confirm fallback rationale is present
    assert "Recommended to reduce key uncertainty before experimental follow-up." in combined_text
    
    forbidden_strings = [
        "[{",
        "}]",
        '"compound_name"',
        '"score_components"',
        '"descriptors"'
    ]
    for forbidden in forbidden_strings:
        assert forbidden not in combined_text


def test_disease_to_lead_workflow_all_cases(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    
    # CASE 1: non-small cell lung cancer / EGFR / Erlotinib
    response = client.post(
        "/api/final-report/create",
        json={
            "report_title": "Case 1",
            "report_mode": "concise_disease_to_lead_report",
            "disease_name": "non-small cell lung cancer",
            "user_entered_target": "EGFR",
            "resolved_target": "Epidermal growth factor receptor",
            "known_compound": "Erlotinib",
            "formats": ["json", "pdf", "docx"],
        }
    )
    assert response.status_code == 200
    res_data = response.json()
    assert res_data["report_id"]
    
    dl_json = client.get(f"/api/final-report/reports/{res_data['report_id']}/json").json()
    assert dl_json["disease_area"] == "non-small cell lung cancer"
    assert dl_json["project_title"] == "Disease-to-Lead: non-small cell lung cancer / EGFR"
    assert dl_json["workflow_input_table"]["disease"] == "non-small cell lung cancer"
    assert dl_json["workflow_input_table"]["user_entered_target"] == "EGFR"
    assert dl_json["workflow_input_table"]["resolved_target"] == "Epidermal growth factor receptor"
    assert dl_json["workflow_input_table"]["target_resolution_status"] == "Synonym Match"
    assert "breast cancer" not in dl_json["executive_summary_paragraph"].lower()
    assert "Resolved target expands or matches the user-entered target." in dl_json["executive_summary_paragraph"]
    assert not any("Resolved target differs" in w for w in dl_json["warnings"])
    
    # CASE 2: breast cancer / HER2 / Lapatinib
    response2 = client.post(
        "/api/final-report/create",
        json={
            "report_title": "Case 2",
            "report_mode": "concise_disease_to_lead_report",
            "disease_name": "breast cancer",
            "user_entered_target": "HER2",
            "resolved_target": "ERBB2",
            "known_compound": "Lapatinib",
            "formats": ["json", "pdf", "docx"],
        }
    )
    assert response2.status_code == 200
    dl_json2 = client.get(f"/api/final-report/reports/{response2.json()['report_id']}/json").json()
    assert dl_json2["disease_area"] == "breast cancer"
    assert dl_json2["workflow_input_table"]["user_entered_target"] == "HER2"
    assert dl_json2["workflow_input_table"]["resolved_target"] == "ERBB2"
    assert dl_json2["workflow_input_table"]["target_resolution_status"] == "Synonym Match"
    assert "egfr" not in dl_json2["executive_summary_paragraph"].lower()
    assert not any("Resolved target differs" in w for w in dl_json2["warnings"])

    # CASE 3: type 2 diabetes / DPP4 / Sitagliptin
    response3 = client.post(
        "/api/final-report/create",
        json={
            "report_title": "Case 3",
            "report_mode": "concise_disease_to_lead_report",
            "disease_name": "type 2 diabetes",
            "user_entered_target": "DPP4",
            "resolved_target": "Dipeptidyl peptidase 4",
            "known_compound": "Sitagliptin",
            "formats": ["json", "pdf", "docx"],
        }
    )
    assert response3.status_code == 200
    dl_json3 = client.get(f"/api/final-report/reports/{response3.json()['report_id']}/json").json()
    assert dl_json3["disease_area"] == "type 2 diabetes"
    assert dl_json3["workflow_input_table"]["user_entered_target"] == "DPP4"
    assert dl_json3["workflow_input_table"]["resolved_target"] == "Dipeptidyl peptidase 4"
    assert dl_json3["workflow_input_table"]["target_resolution_status"] == "Synonym Match"
    assert "cancer" not in dl_json3["executive_summary_paragraph"].lower()
    assert not any("Resolved target differs" in w for w in dl_json3["warnings"])

    # CASE 4: Alzheimer disease / AChE / Donepezil
    response4 = client.post(
        "/api/final-report/create",
        json={
            "report_title": "Case 4",
            "report_mode": "concise_disease_to_lead_report",
            "disease_name": "Alzheimer disease",
            "user_entered_target": "AChE",
            "resolved_target": "Acetylcholinesterase",
            "known_compound": "Donepezil",
            "formats": ["json", "pdf", "docx"],
        }
    )
    assert response4.status_code == 200
    dl_json4 = client.get(f"/api/final-report/reports/{response4.json()['report_id']}/json").json()
    assert dl_json4["disease_area"] == "Alzheimer disease"
    assert dl_json4["workflow_input_table"]["user_entered_target"] == "AChE"
    assert dl_json4["workflow_input_table"]["resolved_target"] == "Acetylcholinesterase"
    assert dl_json4["workflow_input_table"]["target_resolution_status"] == "Synonym Match"
    assert not any("Resolved target differs" in w for w in dl_json4["warnings"])

    # CASE 5: True mismatch test
    response5 = client.post(
        "/api/final-report/create",
        json={
            "report_title": "Case 5",
            "report_mode": "concise_disease_to_lead_report",
            "disease_name": "breast cancer",
            "user_entered_target": "EGFR",
            "resolved_target": "Breast cancer type 1 susceptibility protein",
            "formats": ["json", "pdf", "docx"],
        }
    )
    assert response5.status_code == 200
    dl_json5 = client.get(f"/api/final-report/reports/{response5.json()['report_id']}/json").json()
    assert any("Resolved target differs from user-entered target" in w for w in dl_json5["warnings"])
    assert "WARNING: Resolved target 'Breast cancer type 1 susceptibility protein' differs from user-entered target 'EGFR'" in dl_json5["executive_summary_paragraph"]

    # CASE 6: Repeated-run stale metadata test
    report_a = client.post(
        "/api/final-report/create",
        json={
            "report_title": "Report A",
            "report_mode": "concise_disease_to_lead_report",
            "disease_name": "breast cancer",
            "user_entered_target": "HER2",
            "resolved_target": "ERBB2",
            "known_compound": "Lapatinib",
            "formats": ["json", "pdf", "docx"],
        }
    ).json()
    
    report_b = client.post(
        "/api/final-report/create",
        json={
            "report_title": "Report B",
            "report_mode": "concise_disease_to_lead_report",
            "disease_name": "non-small cell lung cancer",
            "user_entered_target": "EGFR",
            "resolved_target": "Epidermal growth factor receptor",
            "known_compound": "Erlotinib",
            "formats": ["json", "pdf", "docx"],
        }
    ).json()
    
    dl_b = client.get(f"/api/final-report/reports/{report_b['report_id']}/json").json()
    assert dl_b["disease_area"] == "non-small cell lung cancer"
    assert dl_b["project_title"] == "Disease-to-Lead: non-small cell lung cancer / EGFR"
    assert dl_b["workflow_input_table"]["disease"] == "non-small cell lung cancer"
    assert dl_b["workflow_input_table"]["user_entered_target"] == "EGFR"
    assert dl_b["workflow_input_table"]["resolved_target"] == "Epidermal growth factor receptor"
    assert "breast cancer" not in dl_b["executive_summary_paragraph"].lower()
    assert "her2" not in dl_b["executive_summary_paragraph"].lower()
    assert "lapatinib" not in dl_b["executive_summary_paragraph"].lower()


def test_disease_to_lead_report_uses_current_run_candidate_snapshot(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    project = client.post(
        "/api/projects/create",
        json={
            "title": "Shared Disease-to-Lead Project",
            "project_type": "disease_screening",
            "status": "active",
            "disease_area": "type 2 diabetes",
            "target_name": "DPP4",
        },
    ).json()

    _save_disease_to_lead_snapshot(
        project["id"],
        "type 2 diabetes",
        "DPP4",
        "Sitagliptin",
        [
            {
                "compound_name": "TALABOSTAT",
                "compound_id": "CHEMBL93208",
                "smiles": "CC(C)C(N)C(=O)N1CC(O)C1",
                "total_score": 88.0,
                "priority_label": "high_priority_for_review",
                "positive_factors": ["DPP4 snapshot candidate"],
                "risk_factors": [],
                "missing_evidence": [],
            }
        ],
    )
    egfr_run_id = _save_disease_to_lead_snapshot(
        project["id"],
        "breast cancer",
        "EGFR",
        "Erlotinib",
        [
            {
                "compound_name": "Erlotinib",
                "compound_id": "CHEMBL553",
                "smiles": "COCCOc1cc2ncnc(Nc3cccc(Cl)c3)c2cc1OCCOC",
                "total_score": 91.0,
                "priority_label": "high_priority_for_review",
                "positive_factors": ["EGFR current-run candidate"],
                "risk_factors": [],
                "missing_evidence": [],
            }
        ],
    )

    response = client.post(
        "/api/final-report/create",
        json={
            "project_id": project["id"],
            "report_title": "EGFR Current Run Report",
            "report_mode": "concise_disease_to_lead_report",
            "disease_to_lead_run_id": egfr_run_id,
            "disease_name": "breast cancer",
            "user_entered_target": "EGFR",
            "resolved_target": "EGFR",
            "known_compound": "Erlotinib",
            "formats": ["json"],
        },
    )
    assert response.status_code == 200
    report = client.get(response.json()["generated_files"]["json"]).json()
    serialized = json.dumps(report)
    assert "Erlotinib" in serialized
    assert "CHEMBL553" in serialized
    assert "TALABOSTAT" not in serialized
    assert "CHEMBL93208" not in serialized
    assert report["workflow_input_table"]["user_entered_target"] == "EGFR"
    assert report["top_candidate_table"][0]["compound_name"] == "Erlotinib"


def test_final_report_includes_available_trained_model_evidence(tmp_path, monkeypatch):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    monkeypatch.setattr(
        final_report_service,
        "get_active_trained_model_info",
        lambda: {
            "status": "available",
            "model_id": "ames_model_1",
            "model_name": "Ames RF Model",
            "task_name": "AMES",
            "task_type": "binary_classification",
        },
    )
    project = client.post(
        "/api/projects/create",
        json={"title": "Model Evidence Project", "project_type": "disease_screening", "status": "active"},
    ).json()
    run_id = _save_disease_to_lead_snapshot(
        project["id"],
        "breast cancer",
        "EGFR",
        "Erlotinib",
        [
            {
                "compound_name": "Erlotinib",
                "compound_id": "CHEMBL553",
                "smiles": "COCCOc1cc2ncnc(Nc3cccc(Cl)c3)c2cc1OCCOC",
                "total_score": 91.0,
                "priority_label": "high_priority_for_review",
                "trained_model_prediction": {
                    "model_available": True,
                    "active_model_id": "ames_model_1",
                    "model_name": "Ames RF Model",
                    "endpoint_predicted": "AMES",
                    "prediction_label": "inactive",
                    "confidence_level": "High",
                    "uncertainty_score": 0.1,
                    "applicability_domain_status": "inside_domain",
                    "external_validation_status": "not_validated",
                    "evidence_strength": "strong_model_evidence",
                    "missing_evidence": ["external validation"],
                },
                "positive_factors": ["Active trained-model prediction was available."],
                "risk_factors": [],
                "missing_evidence": ["external validation"],
            }
        ],
    )

    response = client.post(
        "/api/final-report/create",
        json={
            "project_id": project["id"],
            "report_mode": "concise_disease_to_lead_report",
            "disease_to_lead_run_id": run_id,
            "disease_name": "breast cancer",
            "user_entered_target": "EGFR",
            "resolved_target": "EGFR",
            "known_compound": "Erlotinib",
            "formats": ["json"],
        },
    )
    assert response.status_code == 200
    report = client.get(response.json()["generated_files"]["json"]).json()
    assert report["has_active_model"] is True
    assert report["model_evidence"]["model_id"] == "ames_model_1"
    assert report["model_evidence"]["model_name"] == "Ames RF Model"
    assert report["model_evidence"]["task_name"] == "AMES"
    assert report["model_evidence"]["evidence_source"] == "trained local model"
    evidence = report["model_evidence_list"][0]
    assert evidence["active_model_id"] == "ames_model_1"
    assert evidence["model_evidence_source"] == "trained local model"
    assert evidence["endpoint_predicted"] == "AMES"
    assert evidence["prediction"] == "inactive"
    assert evidence["applicability_domain_status"] == "inside_domain"


def test_final_report_includes_external_validation_metrics(monkeypatch, tmp_path):
    monkeypatch.setattr(final_report_service, "REPORT_DIR", tmp_path / "final_project_reports")
    monkeypatch.setattr(
        final_report_service,
        "get_active_trained_model_info",
        lambda: {
            "status": "available",
            "model_id": "tox_model_1",
            "model_name": "Toxicity RF Model",
            "task_name": "toxicity_concern",
            "task_type": "binary_classification",
            "model_type": "random_forest",
            "version": "0.18.0-test",
        },
    )
    init_db()
    with get_connection() as connection:
        # Older validation history should remain available elsewhere, but not duplicate the final report summary.
        connection.execute(
            """
            INSERT INTO admet_external_validation_runs (
                model_id, training_run_id, external_dataset_id, task_name, task_type, status,
                valid_count, invalid_count, metric_summary_json, calibration_summary_json, warnings_json, notes
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                "tox_model_1",
                1,
                999,
                "toxicity_concern",
                "binary_classification",
                "completed",
                24,
                0,
                json.dumps({"accuracy": 0.6, "balanced_accuracy": 0.55, "precision": 0.5, "recall": 0.4, "specificity": 0.7, "f1": 0.45, "roc_auc": 0.62}),
                json.dumps({"calibration_status": "available", "calibration_quality": "calibration_poor", "expected_calibration_error": 0.22, "brier_score": 0.31}),
                json.dumps(["Older validation run."]),
                "older external validation test",
            ),
        )
        connection.execute(
            """
            INSERT INTO admet_external_validation_runs (
                model_id, training_run_id, external_dataset_id, task_name, task_type, status,
                valid_count, invalid_count, metric_summary_json, calibration_summary_json, warnings_json, notes
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                "tox_model_1",
                1,
                999,
                "toxicity_concern",
                "binary_classification",
                "completed",
                24,
                0,
                json.dumps({"accuracy": 0.8, "balanced_accuracy": 0.75, "precision": 0.7, "recall": 0.6, "specificity": 0.9, "f1": 0.65, "roc_auc": 0.82}),
                json.dumps({"calibration_status": "available", "calibration_quality": "calibration_moderate", "expected_calibration_error": 0.08, "brier_score": 0.19}),
                json.dumps(["Calibration is dataset-dependent."]),
                "latest external validation test",
            ),
        )
    project = client.post("/api/projects/create", json={"title": "External Validation Report", "project_type": "disease_screening", "status": "active"}).json()
    run_id = _save_disease_to_lead_snapshot(project["id"], "lung cancer", "EGFR", "Erlotinib", [])
    response = client.post(
        "/api/final-report/create",
        json={
            "project_id": project["id"],
            "report_mode": "concise_disease_to_lead_report",
            "disease_to_lead_run_id": run_id,
            "formats": ["json"],
        },
    )
    assert response.status_code == 200
    report = client.get(response.json()["generated_files"]["json"]).json()
    assert report["external_validation"]["has_external_validation"] is True
    assert len(report["external_validation"]["validation_details"]) == 1
    detail = report["external_validation"]["validation_details"][0]
    assert detail["metric_summary"]["accuracy"] == 0.8
    assert detail["notes"] == "latest external validation test"
    assert detail["calibration_summary"]["calibration_quality"] == "calibration_moderate"
